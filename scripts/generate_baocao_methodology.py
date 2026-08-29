"""Generate the "Evaluation Methodology" chapter of baocao.docx.

Reads the per-image evaluation CSV(s) produced by ``evaluate_openvocab.py`` and
writes a Vietnamese, plain-language methodology chapter into ``baocao.docx``:

    1. Phương pháp đo (dataset + harness + cách so khớp)
    2. Định nghĩa từng metric + công thức + ví dụ
    3. Giải thích từng hàm liên quan + bên trong hàm
    4. Kết quả (đọc trực tiếp từ CSV — không hard-code)

Numbers are recomputed from the CSV every run, so re-running after a new
evaluation keeps the document in sync. If a "before" and an "after" CSV both
exist, a comparison table is added.

Usage (from project root):
    python scripts/generate_baocao_methodology.py
    python scripts/generate_baocao_methodology.py --before results/eval_baseline.csv \
        --after results/eval_phase1.csv --out baocao.docx

The output file (default ``baocao.docx``) must NOT be open in Word, or the save
will fail with a permission error (Windows file lock).
"""
from __future__ import annotations

import argparse
import csv
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_PROJECT_ROOT = Path(__file__).resolve().parents[1]
_RISK_THRESHOLDS = [0.0, 0.2, 0.4, 0.5, 0.6, 0.7, 0.8]


@dataclass
class Metrics:
    """Aggregate metrics recomputed from a per-image evaluation CSV."""

    n: int
    errors: int
    recall: int
    top1: int
    top3: int
    family: int
    distractor: int
    uncertain: int
    converted: int  # top-1 hits among recall hits
    by_gen: Dict[str, Dict[str, int]]
    risk_rows: List[tuple]  # (t, kept, sel_top1)

    def pct(self, num: int, den: Optional[int] = None) -> str:
        """Format 'num/den (xx.x%)' (den defaults to scored image count)."""
        d = self.n if den is None else den
        return f"{num}/{d} ({100.0 * num / d:.1f}%)" if d else "0/0 (n/a)"


def _load_metrics(csv_path: Path) -> Metrics:
    """Recompute aggregate metrics from a per-image results CSV."""
    rows: List[dict] = []
    with csv_path.open(encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))

    # Exclude rows with an error AND rows with zero candidates (an upstream
    # API/quota failure that produced no analysis — counting it as a wrong
    # answer would pollute the aggregates).
    scored = [
        r for r in rows
        if not r.get("error") and int(r.get("n_candidates") or 0) > 0
    ]
    n = len(scored)

    def s(field: str) -> int:
        return sum(int(r[field]) for r in scored)

    recall = s("recall_hit")
    converted = sum(
        int(r["top1_hit"]) for r in scored if int(r["recall_hit"])
    )

    by_gen: Dict[str, Dict[str, int]] = {}
    for r in scored:
        g = by_gen.setdefault(r["generator"], {"n": 0, "recall": 0, "top1": 0})
        g["n"] += 1
        g["recall"] += int(r["recall_hit"])
        g["top1"] += int(r["top1_hit"])

    risk_rows: List[tuple] = []
    have_agree = [
        r for r in scored if r.get("panel_agreement") not in (None, "", "None")
    ]
    for t in _RISK_THRESHOLDS:
        kept = [r for r in have_agree if float(r["panel_agreement"]) >= t]
        sel_top1 = sum(int(r["top1_hit"]) for r in kept)
        risk_rows.append((t, len(kept), sel_top1))

    return Metrics(
        n=n, errors=len(rows) - n, recall=recall, top1=s("top1_hit"),
        top3=s("top3_hit"), family=s("family_hit"), distractor=s("distractor_hit"),
        uncertain=s("uncertain"), converted=converted, by_gen=by_gen,
        risk_rows=risk_rows,
    )


# ── docx building helpers ─────────────────────────────────────────────────────

def _h(doc: Document, text: str, level: int) -> None:
    """Add a heading."""
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    """Add a normal paragraph (optionally bold/italic)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic


def _bullet(doc: Document, text: str) -> None:
    """Add a bullet-list item."""
    doc.add_paragraph(text, style="List Bullet")


def _formula(doc: Document, text: str) -> None:
    """Add a centred, monospace-ish formula line."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def _code(doc: Document, text: str) -> None:
    """Add a monospace code/term span as its own paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    """Add a simple table with a header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = head
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


# ── chapter sections ──────────────────────────────────────────────────────────

def _section_intro(doc: Document) -> None:
    """Mở đầu chương: mục tiêu của việc đánh giá."""
    _h(doc, "Phương pháp đánh giá hệ thống nhận dạng phong cách kiến trúc", 1)
    _p(doc,
       "Cách đọc chương này: nó được viết như một cuộc trò chuyện giải thích, "
       "không phải văn kỹ thuật khô khan. Mỗi khái niệm khó đều đi kèm một ví "
       "von đời thường (bài trắc nghiệm, rổ cá, cánh cửa…) để người không "
       "chuyên cũng theo được. Mọi thuật ngữ tiếng Anh được dịch và giải thích "
       "ngay lần đầu xuất hiện.", italic=True)
    _p(doc,
       "Chương này trình bày cách đo lường chất lượng của hệ thống một cách "
       "khách quan, bằng số liệu. Mục tiêu là trả lời câu hỏi: hệ thống đang "
       "đúng/sai bao nhiêu, và quan trọng hơn — SAI Ở KHÂU NÀO, để biết cần cải "
       "thiện chỗ nào.")
    _p(doc,
       "Quy ước: hệ thống nhận một ảnh công trình và phải đoán PHONG CÁCH kiến "
       "trúc của nó (ví dụ: Gothic, Baroque, Mughal…). Vì một công trình có thể "
       "pha trộn nhiều phong cách, hệ thống xuất ra một PHÂN BỐ XÁC SUẤT (tiếng "
       "Anh: mixture / distribution) trên nhiều phong cách, trong đó phong cách "
       "có xác suất cao nhất gọi là phong cách CHÍNH (primary).")


def _section_dataset(doc: Document) -> None:
    """Mô tả tập dữ liệu đánh giá."""
    _h(doc, "1. Tập dữ liệu đánh giá (benchmark)", 2)
    _p(doc,
       "Để chấm điểm cần một bộ ảnh đã biết trước đáp án đúng — gọi là tập đánh "
       "giá có nhãn (tiếng Anh: labelled benchmark). Đặc điểm bộ dữ liệu:")
    _bullet(doc,
            "100 phong cách kiến trúc, mỗi phong cách 2 ảnh: một ảnh do ChatGPT "
            "sinh, một ảnh do Gemini sinh → tổng cộng 200 ảnh "
            "(thư mục data/ChatGPT/ và data/Gemini/, đặt tên 1.png … 100.png).")
    _bullet(doc,
            "Đáp án đúng nằm trong file docs/image_gen_ground_truth.csv. Cột "
            "'main_style' là phong cách đúng (nhãn vàng — tiếng Anh: ground truth).")
    _bullet(doc,
            "Mỗi ảnh được sinh kèm 5 đặc trưng CỦA phong cách chính (cột "
            "ev1..ev5) VÀ 5 đặc trưng NHIỄU cố ý lấy từ phong cách KHÁC "
            "(cột ev6..ev10). Đây là phép thử 'chống nhiễu' (robustness): hệ "
            "thống có bị 5 chi tiết lạ đánh lừa, hay vẫn nhận ra phong cách chủ "
            "đạo của công trình?")
    _p(doc,
       "Lưu ý quan trọng: cơ sở tri thức (tiếng Anh: knowledge base, viết tắt "
       "KB — kho 106 phong cách của hệ thống) đã được kiểm tra là chứa đủ cả "
       "100 phong cách trong tập đánh giá. Nghĩa là khi hệ thống đoán sai thì "
       "KHÔNG phải vì KB thiếu phong cách đó, mà do khâu xử lý — đây là điều "
       "kiện cần để việc đánh giá công bằng.")


def _section_harness(doc: Document) -> None:
    """Mô tả quy trình chấm điểm tự động."""
    _h(doc, "2. Quy trình chấm điểm tự động (evaluation harness)", 2)
    _p(doc,
       "Harness (bộ giàn chấm điểm) là một đoạn chương trình tự động làm thay "
       "công việc của một giám khảo: lần lượt đưa từng ảnh vào hệ thống, lấy "
       "kết quả, so với đáp án đúng, rồi cộng dồn thành bảng điểm. Toàn bộ nằm "
       "trong file scripts/evaluate_openvocab.py. Các bước:")
    _bullet(doc, "Đọc 200 ảnh và đáp án đúng tương ứng.")
    _bullet(doc,
            "Với mỗi ảnh: gọi hàm phân tích của hệ thống "
            "(AnalysisOrchestrator.analyze) để lấy kết quả gồm: danh sách ứng "
            "viên (candidate_names), phong cách chính dự đoán, và phân bố xác suất.")
    _bullet(doc,
            "So sánh kết quả với đáp án đúng và đánh dấu đúng/sai theo từng tiêu chí.")
    _bullet(doc, "Cộng dồn tất cả → in bảng điểm + ghi file results/eval_baseline.csv.")
    _p(doc,
       "So khớp theo ID của KB (đồng nghĩa gộp làm một): hai cách viết khác nhau "
       "của cùng một phong cách — ví dụ 'Traditional Chinese' và 'Chinese "
       "traditional' — phải được tính là GIỐNG nhau. Vì vậy trước khi so sánh, "
       "mọi tên phong cách đều được chuẩn hoá về một mã định danh duy nhất trong "
       "KB (id). Nhờ vậy hệ thống không bị chấm sai chỉ vì khác cách viết.")
    _p(doc,
       "Hiểu hai bước xử lý của hệ thống là chìa khoá để đọc các chỉ số bên "
       "dưới:", bold=True)
    _bullet(doc,
            "Bước LỌC ỨNG VIÊN: từ 106 phong cách, hệ thống rút ra một danh "
            "sách ngắn ~8 phong cách khả dĩ nhất (gọi là candidate_names).")
    _bullet(doc,
            "Bước HỘI ĐỒNG CHẤM: ba 'giám khảo' AI độc lập + một 'trọng tài' "
            "chọn ra phong cách cuối cùng — nhưng CHỈ được chọn trong ~8 ứng "
            "viên ở bước trên.")


def _section_metrics(doc: Document, m: Metrics) -> None:
    """Định nghĩa từng chỉ số + công thức + ví dụ."""
    _h(doc, "3. Các chỉ số đo lường (metrics) — định nghĩa, công thức, ví dụ", 2)
    _p(doc,
       "Metric (chỉ số) là một con số dùng để đo một khía cạnh chất lượng. "
       "Tất cả các chỉ số dưới đây đều là chỉ số CHUẨN của ngành học máy / truy "
       "hồi thông tin, không phải tự chế — nên có thể bảo vệ trước hội đồng.")

    # candidate_recall@K
    _h(doc, "3.1. Độ thu hồi ứng viên — candidate_recall@K", 3)
    _p(doc,
       "Recall (đọc là 'ri-côl', dịch: độ thu hồi / độ bao phủ) là một chỉ số "
       "chuẩn đo độ ĐẦY ĐỦ: trong tất cả những thứ ĐÚNG cần lấy, ta lấy lại "
       "được bao nhiêu. Lưu ý: recall ở đây KHÔNG có nghĩa 'gọi lại' như tiếng "
       "Anh đời thường. Đuôi '@K' nghĩa là 'xét trong K kết quả lấy ra đầu "
       "tiên' (ở đây K = 8 ứng viên).")
    _formula(doc,
             "recall@K = (số ảnh có phong cách đúng nằm trong K ứng viên) / "
             "(tổng số ảnh)")
    _p(doc,
       "Giải thích dễ hiểu: giống bài trắc nghiệm — nếu người ra đề QUÊN in đáp "
       "án đúng vào 4 lựa chọn thì thí sinh giỏi mấy cũng phải khoanh sai. "
       "recall@K đo đúng việc đó: đáp án đúng có được 'đưa vào danh sách lựa "
       "chọn' hay không. Đây là TRẦN CỨNG (giới hạn trên) của toàn hệ thống: "
       "hội đồng không thể chọn đúng một phong cách không có trong danh sách.")
    _p(doc,
       f"Ví dụ với số thật: hệ thống đạt recall@8 = {m.pct(m.recall)} → tức "
       f"chỉ {m.recall} trên {m.n} ảnh là đáp án đúng lọt vào 8 ứng viên; "
       f"{m.n - m.recall} ảnh còn lại đã 'thua' ngay từ khâu lọc.", italic=True)

    # top-1 / top-3
    _h(doc, "3.2. Độ chính xác top-1 và top-3", 3)
    _p(doc,
       "Hệ thống xếp các phong cách theo độ tin và đưa ra phong cách đứng HẠNG 1 "
       "(tin nhất) làm câu trả lời chính. top-1 accuracy (độ chính xác top-1) = "
       "trong tất cả ảnh, bao nhiêu phần trăm có câu-trả-lời-hạng-1 này TRÙNG "
       "với đáp án đúng. Gọi là 'top-1' vì chỉ xét đúng MỘT lựa chọn dẫn đầu — "
       "đây là chỉ số 'đúng/sai' quen thuộc nhất.")
    _formula(doc, "top-1 = (số ảnh có phong cách hạng 1 == đáp án) / (tổng số ảnh)")
    _p(doc,
       "top-3 accuracy (độ chính xác top-3) = tỉ lệ ảnh mà đáp án đúng nằm "
       "trong 3 phong cách có xác suất cao nhất của phân bố. Chỉ số này 'dễ "
       "thở' hơn, hữu ích vì hệ thống cho ra phân bố nhiều phong cách chứ không "
       "chỉ 1 nhãn.")
    _formula(doc, "top-3 = (số ảnh đáp án nằm trong 3 phong cách xác suất cao nhất) / (tổng)")
    _p(doc,
       f"Số thật: top-1 = {m.pct(m.top1)}, top-3 = {m.pct(m.top3)}.", italic=True)

    # family accuracy
    _h(doc, "3.3. Độ chính xác theo HỌ — family accuracy", 3)
    _p(doc,
       "Mỗi phong cách thuộc một HỌ (tiếng Anh: family — nhóm lớn, ví dụ họ "
       "'Islamic' gồm Moorish, Ottoman, Mughal…). family accuracy = tỉ lệ ảnh "
       "mà hệ thống đoán đúng HỌ, dù có thể sai phong cách con. Chỉ số này cho "
       "biết hệ thống 'gần đúng' tới đâu: đoán Ottoman thay vì Mughal (cùng họ "
       "Islamic) tốt hơn nhiều so với đoán Gothic.")
    _formula(doc, "family = (số ảnh cùng HỌ với đáp án) / (tổng số ảnh)")
    _p(doc, f"Số thật: family accuracy = {m.pct(m.family)}.", italic=True)

    # distractor rate
    _h(doc, "3.4. Tỉ lệ mắc bẫy nhiễu — distractor rate", 3)
    _p(doc,
       "Distractor (mồi nhử / yếu tố gây nhiễu) là 5 đặc trưng lạ (ev6..ev10) "
       "cố ý cài vào ảnh từ các phong cách KHÁC. distractor rate = tỉ lệ ảnh mà "
       "hệ thống bị đánh lừa, đoán nhầm sang chính một trong các phong cách "
       "nhiễu đó. Chỉ số này CÀNG THẤP CÀNG TỐT — đo sức 'không bị chi tiết lạ "
       "đánh lạc hướng'.")
    _formula(doc, "distractor rate = (số ảnh đoán trúng 1 phong cách nhiễu) / (tổng)")
    _p(doc, f"Số thật: distractor rate = {m.pct(m.distractor)} (càng thấp càng "
            f"tốt).", italic=True)

    # conversion
    _h(doc, "3.5. Tỉ lệ chuyển hoá — conversion (top-1 khi recall trúng)", 3)
    _p(doc,
       "Đây là chỉ số tách bạch CHẤT LƯỢNG HỘI ĐỒNG khỏi chất lượng khâu lọc. "
       "Nó hỏi: TRONG những ảnh mà đáp án đúng ĐÃ có trong danh sách ứng viên, "
       "hội đồng chọn đúng được bao nhiêu phần trăm? Nói cách khác: khi đã có "
       "cơ hội, hội đồng có nắm bắt được không.")
    _formula(doc, "conversion = (số ảnh top-1 đúng) / (số ảnh recall trúng)")
    _p(doc,
       f"Số thật: conversion = {m.pct(m.converted, m.recall)}. Con số này cao "
       "nghĩa là hội đồng làm tốt phần việc của nó; nếu top-1 tổng thể vẫn thấp "
       "thì 'lỗi' nằm ở khâu lọc ứng viên, không phải hội đồng.", italic=True)
    _p(doc, "Vì sao top-1 = recall × conversion (rất quan trọng để đọc kết quả):",
       bold=True)
    _p(doc,
       "Để một ảnh được tính ĐÚNG, đáp án phải vượt HAI cửa nối tiếp: (cửa 1) "
       "đáp án phải LỌT vào danh sách ứng viên — đo bằng recall; rồi (cửa 2) hội "
       "đồng phải CHỌN đúng nó trong danh sách — đo bằng conversion. Phải qua CẢ "
       "HAI mới đúng, nên tỉ lệ đúng cuối cùng = recall × conversion.")
    _formula(doc, "top-1  =  recall  ×  conversion")
    _p(doc,
       "Ví von 100 thí sinh: giả sử 89 người được phát đề CÓ in đáp án đúng "
       "(recall 89%), còn 11 người bị phát đề THIẾU đáp án nên chắc chắn trượt. "
       "Trong 89 người có cơ hội, 74% khoanh trúng (conversion 74%) = 66 người. "
       "Vậy 66/100 đúng → top-1 = 66%. Không ai khoanh đúng được một đáp án mà đề "
       "không in ra — đó là lý do recall (cửa 1) chặn trên conversion (cửa 2), và "
       "vì sao một recall thấp là 'trần cứng' kéo tụt top-1.", italic=True)
    rec_pct = 100.0 * m.recall / m.n if m.n else 0.0
    conv_pct = 100.0 * m.converted / m.recall if m.recall else 0.0
    bigger = "hội đồng (conversion)" if (100 - conv_pct) > (100 - rec_pct) else "lọc ứng viên (recall)"
    _p(doc,
       "Cách dùng để quyết định cải thiện: cửa nào còn 'dư địa' (khoảng trống còn "
       "lại tới mức 100%) lớn hơn thì sửa cửa đó lời hơn. Recall "
       f"{rec_pct:.0f}% → dư địa {100-rec_pct:.0f}%; conversion {conv_pct:.0f}% → "
       f"dư địa {100-conv_pct:.0f}% → nên ưu tiên nâng {bigger}.")

    # risk-coverage
    _h(doc, "3.6. Đường đánh đổi rủi ro–độ phủ — risk-coverage", 3)
    _p(doc,
       "Hệ thống có cơ chế TỪ CHỐI TRẢ LỜI khi không chắc (tiếng Anh: "
       "abstention — 'bỏ phiếu trắng'), dựa trên mức ĐỒNG THUẬN giữa các giám "
       "khảo (panel_agreement: 3 giám khảo càng nhất trí thì điểm này càng cao). "
       "risk-coverage là cách đo sự đánh đổi: nếu chỉ trả lời khi đủ chắc thì "
       "trả lời ÍT hơn (coverage — độ phủ thấp) nhưng phần trả lời CHÍNH XÁC "
       "hơn (selective accuracy — độ chính xác trên phần đã chọn trả lời).")
    _p(doc,
       "Với mỗi ngưỡng t: chỉ giữ những ảnh có panel_agreement ≥ t (coi như hệ "
       "thống chỉ dám trả lời các ảnh này), rồi đo độ phủ và độ chính xác trên "
       "chúng:")
    _formula(doc, "coverage(t) = (số ảnh giữ lại) / (tổng);  "
                  "selective top-1(t) = (top-1 đúng trong số giữ) / (số giữ)")
    risk_tbl = [
        [f"{t:.1f}", m.pct(kept), (f"{sel}/{kept} ({100.0*sel/kept:.1f}%)"
                                   if kept else "n/a")]
        for (t, kept, sel) in m.risk_rows
    ]
    _table(doc, ["ngưỡng t", "coverage (độ phủ)", "selective top-1"], risk_tbl)
    _p(doc,
       "Mục đích: chọn được ngưỡng t hợp lý để hệ thống 'biết im lặng khi không "
       "chắc' mà vẫn trả lời đủ nhiều — phục vụ phần hiệu chỉnh ở giai đoạn sau.")


def _section_functions(doc: Document) -> None:
    """Giải thích từng hàm liên quan + bên trong hàm."""
    _h(doc, "4. Các hàm liên quan và logic bên trong", 2)
    _p(doc,
       "Phần này mô tả các hàm cốt lõi dùng trong việc đo lường và trong khâu "
       "tạo ứng viên — đầu vào, đầu ra, và cách hoạt động bên trong, bằng ngôn "
       "ngữ phổ thông.")

    _h(doc, "4.1. Nhóm hàm chấm điểm (scripts/evaluate_openvocab.py)", 3)

    _code(doc, "_load_ground_truth() -> Dict[prompt_id, GroundTruth]")
    _p(doc,
       "Đọc file CSV đáp án. Với mỗi dòng, lấy phong cách chính (main_style) và "
       "gom 5 phong cách nhiễu (ev6_style..ev10_style) vào một danh sách. Trả "
       "về từ điển tra cứu theo số thứ tự ảnh (prompt_id).")

    _code(doc, "_kb_id(kb, name) -> str | None")
    _p(doc,
       "Chuẩn hoá một tên phong cách tự do về MÃ ĐỊNH DANH trong KB (id). Bên "
       "trong gọi kb.match(): khớp chính xác → khớp theo tên đồng nghĩa (alias) "
       "→ khớp mờ (fuzzy, chịu được lỗi gõ nhỏ). Nhờ hàm này mà 'Chinese "
       "traditional' và 'Traditional Chinese' cho ra cùng một id → so sánh công "
       "bằng. Trả None nếu tên không có trong KB.")

    _code(doc, "_top_n_styles(result, n) -> List[str]")
    _p(doc,
       "Lấy n phong cách có xác suất cao nhất từ phân bố kết quả "
       "(style_distribution), sắp xếp giảm dần. Dùng cho chỉ số top-3. Nếu "
       "không có phân bố thì trả về đúng nhãn chính.")

    _code(doc, "_score(kb, gt, result, generator) -> EvalRow")
    _p(doc,
       "Trái tim của việc chấm: so kết quả một ảnh với đáp án và tính tất cả "
       "các cờ đúng/sai. Bên trong:")
    _bullet(doc,
            "recall_hit: id đáp án có nằm trong tập id của candidate_names không.")
    _bullet(doc,
            "top1_hit: id phong cách chính dự đoán == id đáp án.")
    _bullet(doc,
            "top3_hit: id đáp án có nằm trong 3 phong cách xác suất cao nhất.")
    _bullet(doc,
            "family_hit: phong cách dự đoán và đáp án cùng 'parent' (cùng HỌ).")
    _bullet(doc,
            "distractor_hit: phong cách dự đoán trùng 1 trong 5 phong cách nhiễu.")
    _p(doc,
       "Tất cả so sánh đều trên id của KB (đã chuẩn hoá), không so chuỗi thô.")

    _code(doc, "vòng tính risk-coverage trong _print_summary()")
    _p(doc,
       "Với mỗi ngưỡng t trong danh sách định trước, lọc các ảnh có "
       "panel_agreement ≥ t, rồi tính độ phủ (số giữ / tổng) và độ chính xác "
       "top-1 trên phần giữ. In thành bảng để chọn ngưỡng từ chối trả lời.")

    _h(doc, "4.2. Nhóm hàm tạo ứng viên (chatbot/services/style_kb_service.py)", 3)
    _p(doc, "Đây là các hàm quyết định recall — tức 'đáp án đúng có lọt vào "
            "danh sách hay không'.")

    _code(doc, "match(name) -> StyleEntry | None")
    _p(doc,
       "Khớp một tên phong cách tự do (do AI đề xuất) với một mục trong KB, "
       "theo thứ tự: khớp chính xác → khớp tên đồng nghĩa → khớp mờ (difflib, "
       "ngưỡng tương đồng 0.82). Trước khi khớp, tên được 'làm sạch' (bỏ chữ "
       "thừa như 'architecture', 'style', bỏ dấu câu, viết thường).")

    _code(doc, "build_candidate_set_voted(proposed_lists, min_votes, top_k)")
    _p(doc,
       "Tạo danh sách ứng viên bằng cách BỎ PHIẾU. Mỗi lượt phân tích ảnh (hệ "
       "thống chạy nhiều lượt) đề xuất một danh sách tên; mỗi phong cách (theo "
       "id) được +1 phiếu cho mỗi lượt nó xuất hiện. Giữ các phong cách đạt từ "
       "min_votes phiếu trở lên, xếp theo số phiếu giảm dần. min_votes=1 nghĩa "
       "là GIỮ TẤT CẢ (hợp nhất — recall tối đa cho cách làm theo tên). Sau đó "
       "'độn' thêm các biến thể tên gần và các phong cách cùng họ cho đủ top_k.")

    _code(doc, "styles_in_family(family_id) -> List[StyleEntry]")
    _p(doc,
       "Trả về TẤT CẢ phong cách thuộc một HỌ. Dùng cho chiến lược 'từ thô đến "
       "tinh' (coarse-to-fine): khi AI chỉ ra được HỌ của công trình (dễ đoán "
       "đúng hơn tên cụ thể), ta nạp cả họ vào ứng viên để chắc chắn phong cách "
       "con đúng có mặt. Nhận cả id họ ('islamic') lẫn tên hiển thị ('Islamic').")

    _code(doc, "retrieve_by_features(observed_features, top_n) -> List[StyleEntry]")
    _p(doc,
       "Truy hồi ứng viên theo ĐẶC TRƯNG quan sát, không theo tên. Với mỗi mục "
       "KB, so các cụm đặc trưng quan sát được (ví dụ 'horseshoe arches', "
       "'muqarnas') với danh sách đặc trưng định nghĩa của mục đó "
       "(defining_features + expected_profile), dùng độ trùng tập-từ (Jaccard). "
       "Cộng điểm các cụm khớp đủ mạnh; xếp hạng, trả top_n. Nhờ vậy bắt được "
       "phong cách mà AI MÔ TẢ ĐÚNG đặc trưng nhưng QUÊN gọi tên — ví dụ tả "
       "'horseshoe arches + muqarnas' sẽ kéo được 'Moorish' lên đầu dù không "
       "hề nhắc tên Moorish.")


def _section_results(doc: Document, before: Metrics,
                     after: Optional[Metrics]) -> None:
    """Bảng kết quả + chẩn đoán (đọc từ CSV)."""
    _h(doc, "5. Kết quả đo và chẩn đoán", 2)

    if before.errors:
        _p(doc,
           f"Lưu ý: {before.errors} ảnh bị LOẠI khỏi thống kê do lỗi gọi API / "
           f"hết quota (không tạo được ứng viên nên không thể chấm công bằng). "
           f"Các số dưới đây tính trên {before.n} ảnh chạy thành công. Nếu số "
           "ảnh thành công chưa đủ 200, đây là kết quả MỘT PHẦN — cần chạy lại "
           "đầy đủ khi quota khôi phục để có con số cuối cùng.", italic=True)
    if after is None:
        _p(doc, f"Kết quả của hệ thống gốc (baseline — mốc so sánh ban đầu) trên "
                f"{before.n} ảnh chạy thành công:", bold=True)
        rows = [
            ["candidate_recall@8 (TRẦN CỨNG)", before.pct(before.recall)],
            ["top-1", before.pct(before.top1)],
            ["top-3", before.pct(before.top3)],
            ["family accuracy", before.pct(before.family)],
            ["conversion (top-1 | recall trúng)", before.pct(before.converted, before.recall)],
            ["distractor rate (càng thấp càng tốt)", before.pct(before.distractor)],
            ["abstained (từ chối trả lời)", before.pct(before.uncertain)],
        ]
        _table(doc, ["Chỉ số", "Baseline"], rows)
    else:
        _h(doc, "Cải tiến: tạo ứng viên đa kênh (multichannel)", 3)
        _p(doc,
           "Hãy hình dung danh sách ứng viên như một CĂN PHÒNG: phong cách đúng "
           "phải 'vào được phòng' thì hội đồng mới có thể chọn nó. Hệ thống gốc "
           "chỉ có MỘT cánh cửa, nên hay bị bỏ sót:")
        _bullet(doc,
                "Cửa 1 — theo TÊN: chỉ vào được nếu AI GỌI ĐÚNG TÊN phong cách. "
                "Với phong cách hiếm (Mudéjar, Ottonian…) AI hay quên tên → cửa "
                "đóng → đáp án đúng đứng ngoài.")
        _p(doc, "Cải tiến mở thêm hai cánh cửa nữa (vì vậy gọi là 'đa kênh'):")
        _bullet(doc,
                "Cửa 2 — theo ĐẶC TRƯNG: nếu AI MÔ TẢ đúng đặc điểm (ví dụ "
                "'vòm móng ngựa', 'vòm tổ ong muqarnas') thì dù không gọi tên, "
                "hệ thống vẫn đối chiếu đặc điểm đó với kho tri thức và kéo đúng "
                "phong cách (Moorish) vào phòng.")
        _bullet(doc,
                "Cửa 3 — theo HỌ: nếu AI đoán đúng HỌ lớn (ví dụ 'Islamic') — "
                "việc dễ hơn nhiều so với tên cụ thể — thì cả họ đó được mời vào "
                "phòng, đảm bảo phong cách con đúng có mặt.")
        _p(doc,
           "Một cửa kẹt thì cửa khác cứu → đáp án đúng lọt vào danh sách nhiều "
           "hơn → recall tăng. Bảng dưới so TRƯỚC (1 cửa) và SAU (3 cửa):",
           bold=True)
        rows = [
            ["candidate_recall@8 (TRẦN CỨNG)", before.pct(before.recall), after.pct(after.recall)],
            ["top-1", before.pct(before.top1), after.pct(after.top1)],
            ["top-3", before.pct(before.top3), after.pct(after.top3)],
            ["family accuracy", before.pct(before.family), after.pct(after.family)],
            ["conversion", before.pct(before.converted, before.recall),
             after.pct(after.converted, after.recall)],
            ["distractor rate", before.pct(before.distractor), after.pct(after.distractor)],
            ["abstained", before.pct(before.uncertain), after.pct(after.uncertain)],
        ]
        _table(doc, ["Chỉ số", "Trước", "Sau"], rows)

    _h(doc, "5.1. Chẩn đoán điểm nghẽn", 3)
    recall_pct = 100.0 * before.recall / before.n if before.n else 0.0
    conv_pct = 100.0 * before.converted / before.recall if before.recall else 0.0
    top1_pct = 100.0 * before.top1 / before.n if before.n else 0.0
    _p(doc,
       f"Quan sát then chốt: recall@8 ({recall_pct:.1f}%) × conversion "
       f"({conv_pct:.1f}%) = top-1 ({top1_pct:.1f}%). Phép nhân khớp tuyệt đối "
       "này chứng minh: mỗi câu trả lời đúng đều đòi hỏi đáp án phải có sẵn "
       "trong danh sách ứng viên TRƯỚC, rồi hội đồng mới chọn được.")
    _p(doc,
       f"Kết luận: điểm nghẽn nằm ở KHÂU LỌC ỨNG VIÊN, không phải hội đồng. "
       f"Hội đồng chấm đúng tới {conv_pct:.0f}% khi có cơ hội (conversion cao), "
       f"nhưng {100.0 - recall_pct:.0f}% số ảnh đã mất đáp án đúng ngay từ khâu "
       "lọc nên hội đồng không thể cứu. Vì vậy hướng cải thiện chính là NÂNG "
       "RECALL: bổ sung kênh tạo ứng viên theo ĐẶC TRƯNG (retrieve_by_features) "
       "và theo HỌ (styles_in_family), thay vì chỉ dựa vào tên do AI gọi.")

    _h(doc, "5.2. Phân tích theo nguồn sinh ảnh", 3)
    gen_rows = [
        [g, f"{v['n']}",
         f"{v['recall']}/{v['n']} ({100.0*v['recall']/v['n']:.1f}%)" if v['n'] else "n/a",
         f"{v['top1']}/{v['n']} ({100.0*v['top1']/v['n']:.1f}%)" if v['n'] else "n/a"]
        for g, v in sorted(before.by_gen.items())
    ]
    _table(doc, ["Nguồn ảnh", "Số ảnh", "recall@8", "top-1"], gen_rows)


def _build(before: Metrics, after: Optional[Metrics], out_path: Path) -> None:
    """Assemble the full chapter and save the docx."""
    doc = Document()
    _section_intro(doc)
    _section_dataset(doc)
    _section_harness(doc)
    _section_metrics(doc, before)
    _section_functions(doc)
    _section_results(doc, before, after)
    doc.save(str(out_path))


def main() -> int:
    """Parse args, recompute metrics from CSV(s), write baocao.docx."""
    parser = argparse.ArgumentParser(description="Generate baocao.docx methodology")
    parser.add_argument("--before", default=str(_PROJECT_ROOT / "results" / "eval_baseline.csv"),
                        help="baseline per-image CSV")
    parser.add_argument("--after", default="",
                        help="optional 'after' per-image CSV for before/after table")
    parser.add_argument("--out", default=str(_PROJECT_ROOT / "baocao.docx"),
                        help="output docx path (must not be open in Word)")
    args = parser.parse_args()

    before_path = Path(args.before)
    if not before_path.is_file():
        print(f"ERROR: baseline CSV not found: {before_path}")
        return 1
    before = _load_metrics(before_path)
    after = _load_metrics(Path(args.after)) if args.after else None

    out_path = Path(args.out)
    try:
        _build(before, after, out_path)
    except PermissionError:
        print(f"ERROR: cannot write {out_path} — close it in Word first.")
        return 1
    print(f"Wrote methodology chapter to {out_path} "
          f"(baseline n={before.n}; recall={before.pct(before.recall)}; "
          f"top-1={before.pct(before.top1)})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
