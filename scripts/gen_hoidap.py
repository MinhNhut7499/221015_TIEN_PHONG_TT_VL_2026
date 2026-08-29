# -*- coding: utf-8 -*-
"""Sinh file hoidap.docx — ghi lại toàn bộ câu hỏi & phản hồi của phiên làm việc.

Chạy: python scripts/gen_hoidap.py
Yêu cầu: pip install python-docx (chỉ để sinh tài liệu, không phải dependency dự án).
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT = Path(__file__).resolve().parent.parent / "hoidap.docx"


def add_question(doc: Document, idx: int, text: str) -> None:
    """Thêm một câu hỏi (in đậm, có nền nhấn)."""
    h = doc.add_heading(level=2)
    run = h.add_run(f"Câu hỏi {idx}")
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)


def add_answer_heading(doc: Document) -> None:
    """Tiêu đề 'Phản hồi'."""
    p = doc.add_paragraph()
    r = p.add_run("Phản hồi")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0x2E)


def para(doc: Document, text: str) -> None:
    """Đoạn văn thường."""
    doc.add_paragraph(text)


def bullet(doc: Document, text: str) -> None:
    """Gạch đầu dòng."""
    doc.add_paragraph(text, style="List Bullet")


def code(doc: Document, text: str) -> None:
    """Khối mã/đồ hoạ dạng monospace."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main() -> None:
    doc = Document()

    title = doc.add_heading("Hỏi – Đáp phiên làm việc: PIVOT Open-Vocabulary (Ngày 2)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Dự án Architecture-AI · Ngày 2026-06-14 · Ghi lại toàn bộ câu hỏi của người "
        "dùng và phản hồi của trợ lý trong phiên."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Q1 ────────────────────────────────────────────────────────────────────
    add_question(doc, 1, "Đọc CLAUDE.md và cho tôi biết: dự án đang ở đâu, bước tiếp theo là gì?")
    add_answer_heading(doc)
    para(doc, "Dự án vừa hoàn thành Ngày 1/5 của PIVOT sang Open-Vocabulary (nhận dạng "
              "phong cách mở, bỏ kiến trúc closed-set 10 lớp cũ). 154 test pass. Đã có:")
    bullet(doc, "KB (cơ sở tri thức) 106 phong cách / 12 họ trong chatbot/knowledge/styles.json.")
    bullet(doc, "StyleKbService — khớp tên phong cách, build tập ứng viên.")
    bullet(doc, "Schemas EvidenceSheet/EvidenceItem/StyleEntry; config mới (STYLE_KB_PATH, RETRIEVAL_TOP_K, PANEL_MODELS).")
    bullet(doc, "Các file cũ (YOLO, ResNet, fuser, CV) cố ý chưa xóa để không vỡ 154 test.")
    para(doc, "Bước tiếp theo (theo thứ tự):")
    bullet(doc, "Ngày 2: Agent A (VLM điền phiếu bằng chứng) + KB grounding + viết lại orchestrator.")
    bullet(doc, "Ngày 3: Hội đồng 3 giám khảo + trọng tài + abstention (từ chối).")
    bullet(doc, "Ngày 4: harness đánh giá trên Kaggle 25 phong cách (vượt baseline Xu 2014 = 46%@25).")
    bullet(doc, "Ngày 5: demo phi-phương-Tây zero-shot + smoke API + cập nhật tài liệu.")

    # ── Q2 ────────────────────────────────────────────────────────────────────
    add_question(doc, 2, "hãy làm tiếp phần việc của ngày 2")
    add_answer_heading(doc)
    para(doc, "Bắt đầu khám phá mã nguồn (giai đoạn lập kế hoạch) — đọc schemas, "
              "pipeline_runner, orchestrator, prompt_builder, style_kb_service, gemini_service "
              "để hiểu hiện trạng trước khi viết kế hoạch chi tiết cho Ngày 2.")

    # ── Q3 ────────────────────────────────────────────────────────────────────
    add_question(doc, 3, "PIVOT là gì? Hãy viết cho tôi sơ đồ giải thuật mới, mã giả giải "
                          "thuật mới và giải thích từng bước.")
    add_answer_heading(doc)
    para(doc, "PIVOT = thay đổi toàn bộ kiến trúc giải thuật từ closed-set (tập đóng) 10 lớp "
              "sang open-vocabulary (từ vựng mở) 100+ phong cách.")
    para(doc, "Hệ thống CŨ (bị thay thế):")
    code(doc, "Ảnh → YOLO detect component → ResNet50 prior (10 lớp) → 7 LLM agent → 1 trong 10 nhãn cố định")
    para(doc, "Vấn đề: train model với 10 phong cách phương Tây, không mở rộng được. Giảng viên "
              "yêu cầu nhận dạng 100+ phong cách toàn cầu (Mughal, Khmer, Ottoman, Maya...).")
    para(doc, "Sơ đồ giải thuật MỚI (6 stage):")
    code(doc,
         "ẢNH ĐẦU VÀO\n"
         "   │\n"
         "   ▼ STAGE A — Agent A (Gemini VLM)\n"
         "   Điền PHIẾU BẰNG CHỨNG 12 chiều (massing/roof/supports/arch/openings/\n"
         "   facade/ornament/material/verticals/vault_dome/spatial_org/diagnostic)\n"
         "   Mỗi chiều: { feature, bbox, suggested_styles }\n"
         "   → proposed_styles = dedup mọi suggested_styles\n"
         "   │\n"
         "   ▼ STAGE B — KB Grounding (mã thuần, KHÔNG LLM)\n"
         "   proposed_styles → khớp 106 mục KB → Top-K ứng viên\n"
         "   tên ngoài-KB → hàng đợi đề xuất\n"
         "   │\n"
         "   ▼ STAGE C — Hội đồng 3 Giám Khảo ĐỘC LẬP  [Ngày 3]\n"
         "   Gemini / DeepSeek / GPT-4o chấm điểm RIÊNG LẺ → 3 phân bố\n"
         "   Đo đồng thuận liên-giám-khảo\n"
         "   │\n"
         "   ▼ STAGE D — Trọng Tài (GPT-4o + ảnh đầy đủ)  [Ngày 3]\n"
         "   Hoà giải, BỊ RÀNG BUỘC bởi bằng chứng + KB\n"
         "   │\n"
         "   ▼ STAGE E — Abstention (Từ chối) nếu đồng thuận/độ chắc thấp\n"
         "   │\n"
         "   ▼ STAGE F — Narrative song ngữ (Anh + Việt), bám từng chiều bằng chứng")
    para(doc, "Mã giả (pseudocode):")
    code(doc,
         "function analyze(image_bytes):\n"
         "  # STAGE A\n"
         "  evidence_sheet = gemini.fill_evidence_sheet(image_bytes)\n"
         "  # STAGE B\n"
         "  candidates, out_of_kb = kb.build_candidate_set(evidence_sheet.proposed_styles, top_k=8)\n"
         "  # STAGE C (Ngày 3)\n"
         "  dist_gemini   = gemini.judge(evidence_sheet, candidates)\n"
         "  dist_deepseek = deepseek.judge(evidence_sheet, candidates)\n"
         "  dist_openai   = openai.judge(evidence_sheet, candidates)\n"
         "  agreement = inter_rater_agreement(dist_gemini, dist_deepseek, dist_openai)\n"
         "  # STAGE E\n"
         "  if agreement < THRESHOLD: return abstain()\n"
         "  # STAGE D\n"
         "  final = gpt4o.arbitrate(image, [3 phân bố], evidence_sheet, candidates)\n"
         "  # STAGE F\n"
         "  return result(final, narrative_en, narrative_vi)")
    para(doc, "Giải thích từng stage: A = VLM là 'thanh tra' điền biên bản có cấu trúc; "
              "B = khớp tên tự do về mục KB chuẩn; C = 3 LLM khác nhà cung cấp chấm độc lập để "
              "phát hiện ảo giác; D = trọng tài hoà giải dựa bằng chứng; E = nói 'không biết' khi "
              "không chắc; F = giải thích song ngữ truy vết được.")

    # ── Q4 ────────────────────────────────────────────────────────────────────
    add_question(doc, 4,
                 "Lưu ý confidence của LLM không đáng tin, phải có hướng giải quyết. "
                 "proposed_styles là gì, dedup là sao? Cho ví dụ chi tiết từng bước. "
                 "exact→alias→fuzzy 0.82 là gì, 0.82 có hardcode? Giải thích rõ Stage B. "
                 "Stage C: giám khảo nhận kết quả Agent A liệu đã đúng chưa? Đầu ra Stage C "
                 "là gì, dist_gemini kiểu dữ liệu gì? Giải thích Krippendorff's alpha và "
                 "pairwise rank correlation. Các agent khác đâu? Đưa quy tắc 'không chêm "
                 "tiếng Anh nếu không giải thích' vào bộ quy tắc.")
    add_answer_heading(doc)
    para(doc, "1) Confidence (độ chắc chắn LLM tự gán) KHÔNG đáng tin. Hướng giải quyết:")
    bullet(doc, "KHÔNG dùng con số confidence thô làm trọng số nhân.")
    bullet(doc, "Nếu buộc dùng → chỉ coi là thứ bậc (cao/trung bình/thấp), không dùng giá trị liên tục.")
    bullet(doc, "Tín hiệu bất định CHÍNH = đồng thuận giữa giám khảo + phương sai khi lấy mẫu lặp (Monte Carlo).")
    para(doc, "2) proposed_styles = gộp tất cả suggested_styles của 12 chiều rồi bỏ trùng (dedup "
              "= deduplicate = loại bản trùng). Đây là danh sách ứng viên đưa vào Stage B.")
    para(doc, "3) Ví dụ Nhà thờ Đức Bà Paris:")
    bullet(doc, "Stage A: chiều arch → feature 'cuốn nhọn', suggested ['Gothic']; chiều supports → "
                "'trụ bổ tường bay', ['Gothic','High Gothic']; chiều material → 'đá vôi', "
                "['Gothic','Romanesque']; diagnostic → 'cuốn nhọn + trụ bay + cửa hoa hồng', ['French Gothic'].")
    bullet(doc, "proposed_styles = ['Gothic','High Gothic','Rayonnant Gothic','Romanesque','French Gothic'].")
    bullet(doc, "Stage B: khớp về KB → candidates [Gothic, Romanesque, Norman, Byzantine, ...] (top-8).")
    bullet(doc, "Stage C: dist_gemini={Gothic:0.88,...}, dist_deepseek={Gothic:0.82,...}, dist_openai={Gothic:0.91,...} → đồng thuận cao.")
    bullet(doc, "Stage D: final={Gothic:0.87, Romanesque:0.08, ...}. Stage E: không từ chối. Stage F: narrative song ngữ.")
    para(doc, "4) exact→alias→fuzzy 0.82 = ba tầng khớp tên, thử từ chặt đến lỏng:")
    bullet(doc, "exact: trùng y hệt tên/mã KB sau chuẩn hoá.")
    bullet(doc, "alias: trùng một bí danh (aliases) của mục KB.")
    bullet(doc, "fuzzy: so độ giống chuỗi bằng difflib; nếu ≥ ngưỡng (0.82) thì khớp. 0.82 vốn là "
                "default trong code → đã đưa lên config thành KB_FUZZY_CUTOFF (bỏ hard-code).")
    para(doc, "5) Stage B là mã thuần, KHÔNG gọi LLM — chỉ tra cứu chuỗi trong RAM; tên ngoài-KB "
              "ghi vào hàng đợi đề xuất (chờ người duyệt), không tự thêm vào KB.")
    para(doc, "6) Stage C: giám khảo KHÔNG nhận PHÁN QUYẾT của Agent A, chỉ nhận QUAN SÁT có khoanh "
              "vùng (kiểm chứng được bằng mắt); trọng tài còn nhìn lại ảnh gốc nên có thể bác lỗi. "
              "Đầu ra Stage C = 3 StyleDistribution (mỗi giám khảo một cái), tức từ điển {tên phong "
              "cách → xác suất} tổng = 1. KHÔNG lấy trung bình ngay; giữ cả 3 + đo độ lệch.")
    bullet(doc, "Krippendorff's alpha = thước đo mức đồng thuận giữa nhiều người chấm, có hiệu chỉnh "
                "may rủi; từ -1 đến 1 (1 = đồng thuận hoàn hảo).")
    bullet(doc, "Pairwise rank correlation = tương quan THỨ HẠNG từng cặp giám khảo (ví dụ Spearman), "
                "không phụ thuộc con số xác suất tuyệt đối — phù hợp với nguyên tắc confidence không đáng tin.")
    para(doc, "7) Danh sách agent hệ MỚI: Agent A (Gemini, điền phiếu); KB grounding (mã thuần); "
              "Giám khảo 1/2/3 (Gemini/DeepSeek/GPT-4o); Trọng tài (GPT-4o + ảnh). Ngày 2 tạm dùng "
              "Agent 5/6/7 cũ làm Tier 2.")
    para(doc, "8) Quy tắc ngôn ngữ đã ghi: KHÔNG bao giờ chêm từ/thuật ngữ tiếng Anh mà không kèm "
              "giải nghĩa tiếng Việt ngay tại chỗ (đã lưu vào memory feedback_language).")
    para(doc, "Người dùng đã chọn (qua bảng hỏi): bỏ HẲN trường confidence; phạm vi Ngày 2 = Stage "
              "A+B + tái dùng Agent 5/6/7; xóa file chết NGAY trong Ngày 2.")

    # ── Q5 ────────────────────────────────────────────────────────────────────
    add_question(doc, 5,
                 "Tại sao phải loại bỏ trùng trong khi có thể dựa vào suggested_styles bị trùng "
                 "để hỗ trợ kết luận (nhiều chiều có Gothic = bằng chứng hỗ trợ)? Ở Stage C, các "
                 "điểm chấm có được dùng làm giá trị cho phần sau không hay chỉ để xếp hạng?")
    add_answer_heading(doc)
    para(doc, "Suy nghĩ của bạn ĐÚNG. Có hai con đường song song, không xung đột:")
    bullet(doc, "(a) dedup → proposed_styles: CHỈ để tra KB chọn ứng viên (tra tên trùng ra cùng mục KB = vô nghĩa).")
    bullet(doc, "(b) đếm tần suất → _build_evidence_votes: 'bao nhiêu chiều ủng hộ phong cách X' = tín hiệu hỗ trợ. "
                "Tần suất KHÔNG bị dedup vứt đi.")
    para(doc, "Cảnh báo: 12 chiều KHÔNG độc lập (hiệu ứng mỏ neo — Agent A nghĩ Gothic thì điền Gothic "
              "nhiều chiều). Nên tần suất chỉ là tín hiệu MỀM để xếp hạng/đưa vào prompt, KHÔNG phải "
              "xác suất cuối; quyết định cuối do giám khảo + trọng tài từ bằng chứng + KB.")
    para(doc, "Stage C — điểm chấm của giám khảo:")
    bullet(doc, "CÓ dùng: (1) đo đồng thuận để quyết định abstain — qua THỨ HẠNG, không qua giá trị tuyệt đối; "
                "(2) làm ngữ cảnh cho trọng tài.")
    bullet(doc, "KHÔNG dùng: lấy trung bình số học 3 phân bố làm kết quả (trung bình che giấu bất đồng). "
                "Xác suất cuối do trọng tài tự chốt, kèm margin/entropy/cờ uncertain.")

    # ── Q6 ────────────────────────────────────────────────────────────────────
    add_question(doc, 6, "vậy hãy dùng pairwise rank correlation cho hệ thống này")
    add_answer_heading(doc)
    para(doc, "Đã chốt: pairwise rank correlation (tương quan thứ hạng Spearman) làm thước đo đồng "
              "thuận Stage C. Với mỗi cặp giám khảo (Gemini–DeepSeek, Gemini–GPT, DeepSeek–GPT) tính "
              "tương quan thứ hạng các phong cách, rồi lấy trung bình. Cách này không phụ thuộc con số "
              "xác suất tuyệt đối (đúng nguyên tắc confidence không đáng tin). Đây là việc Ngày 3 "
              "(Stage C) — đã ghi vào kế hoạch; triển khai bằng scipy.stats.spearmanr.")

    # ── Q7 ────────────────────────────────────────────────────────────────────
    add_question(doc, 7, "hãy làm tiếp phần việc của ngày 2 (triển khai sau khi duyệt kế hoạch)")
    add_answer_heading(doc)
    para(doc, "Đã triển khai xong Ngày 2 — 82 test pass, mã sạch (không còn import torch/yolo/fuser/CV).")
    para(doc, "Tạo mới:")
    bullet(doc, "chatbot/services/evidence_extractor.py — Agent A (Gemini điền phiếu 12 chiều) + MockEvidenceExtractor; "
                "đổi bbox đúng từng trục (x×rộng, y×cao); JSON hỏng → sheet rỗng không văng lỗi.")
    bullet(doc, "chatbot/utils/json_utils.py — parse_json_safe dùng chung.")
    bullet(doc, "test/test_evidence_extractor.py — 7 test.")
    para(doc, "Viết lại:")
    bullet(doc, "pipeline_runner.py — open-vocab thuần: Agent 5 (advocate) → 6 (alternative) → 7 (arbiter, GPT-4o + ảnh). "
                "_build_evidence_votes đếm tần suất, KHÔNG dùng confidence.")
    bullet(doc, "prompt_builder.py — build_agent_a_prompt (không xin confidence) + Agent 5/6/7 ăn evidence sheet + ứng viên KB.")
    bullet(doc, "analysis_orchestrator.py — Stage A (Agent A) → Stage B (KB grounding).")
    bullet(doc, "schemas.py — bỏ STYLE_CLASSES, bỏ confidence khỏi EvidenceItem, PipelineInput mới, bỏ AttributeVector/GlobalFeatureOutput.")
    bullet(doc, "config.py — thêm KB_FUZZY_CUTOFF=0.82, xóa field closed-set.")
    para(doc, "Đã xóa: 8 service + 5 util (fusion/rule_checker/cv_attributes/gradcam/metrics) + 5 script + "
              "4 weights + 4 test file + 9 deps nặng (torch/torchvision/ultralytics/sklearn/opencv/scikit-image/"
              "numpy/imagehash/grad-cam) → máy RTX 3050 hết lo CUDA.")
    para(doc, "Đã ghi nhận: memory feedback_language (quy tắc không chêm tiếng Anh không giải nghĩa); "
              "CLAUDE.md cập nhật Ngày 2/5 + quyết định #73 (bỏ confidence), #74 (xóa ngay → 82 test), "
              "#75 (pairwise rank correlation cho Stage C).")
    para(doc, "Còn lại Ngày 3-5: tách 3 giám khảo độc lập thật + Monte Carlo + đo đồng thuận Spearman + "
              "abstention; harness eval Kaggle 25; provenance KB; smoke end-to-end qua API (chưa chạy, cần quota LLM).")

    doc.save(OUT)
    print(f"Đã tạo: {OUT}")


if __name__ == "__main__":
    main()
