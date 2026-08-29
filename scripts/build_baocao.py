# -*- coding: utf-8 -*-
"""Dựng lại nội dung Báo cáo Khóa luận Tốt nghiệp trên khung định dạng của template.

Mở ``BaoCaoKhoaLuanTotNghiep.template.docx`` (bản sao gốc), giữ nguyên styles +
section + trang bìa, sửa định dạng theo Quy định 309-QĐ ĐHNCT, rồi viết mới toàn
bộ phần đầu + 6 chương + tài liệu tham khảo cho đề tài "Hệ thống tự động nhận dạng
phong cách kiến trúc từ ảnh chụp sử dụng trí tuệ nhân tạo".

Chạy:  python scripts/build_baocao.py
(Đóng Word/Excel trước khi chạy để tránh khóa file.)
"""
from __future__ import annotations

import os
import shutil

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE = os.path.join(ROOT, "BaoCaoKhoaLuanTotNghiep.template.docx")
SOURCE = os.path.join(ROOT, "BaoCaoKhoaLuanTotNghiep.docx")
OUTPUT = os.path.join(ROOT, "BaoCaoKhoaLuanTotNghiep.docx")

FONT = "Times New Roman"
SZ_BODY = Pt(13)
SZ_CHAPTER = Pt(14)


# ───────────────────────────── helpers: low-level XML ──────────────────────────
def _set_font(run, size=SZ_BODY, bold=False, italic=False, name=FONT):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    # ensure East-Asian/complex script also use the font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def _para_format(p, *, line=1.2, before=6, after=0, indent_cm=None,
                 align=None, keep_with_next=False):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent_cm is not None:
        pf.first_line_indent = Cm(indent_cm)
    if align is not None:
        pf.alignment = align
    pf.keep_with_next = keep_with_next


def _new_p_after(ref_p_elem):
    """Tạo <w:p> mới ngay sau phần tử paragraph ref_p_elem, trả về Paragraph."""
    new = OxmlElement("w:p")
    ref_p_elem.addnext(new)
    from docx.text.paragraph import Paragraph
    return Paragraph(new, None)


def _repeat_header(row):
    """Đánh dấu hàng là hàng tiêu đề (lặp lại khi bảng sang trang)."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
    trPr.append(th)


def _move_sectpr(src_p, dst_p):
    """Di chuyển <w:sectPr> từ pPr của src_p sang dst_p (giữ section break)."""
    src_ppr = src_p.get_or_add_pPr()
    sect = src_ppr.find(qn("w:sectPr"))
    if sect is None:
        return
    src_ppr.remove(sect)
    dst_ppr = dst_p.get_or_add_pPr()
    dst_ppr.append(sect)


# ───────────────────────────── helpers: content writers ────────────────────────
class Doc:
    """Bao bọc python-docx Document + bộ đếm hình/bảng theo chương."""

    def __init__(self, document):
        self.d = document
        self.chapter = 0
        self.fig = 0
        self.tbl = 0

    # -- paragraphs --------------------------------------------------------------
    def para(self, text="", *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             italic=False, bold=False, line=1.2, before=6, after=0, size=SZ_BODY):
        p = self.d.add_paragraph(style="Normal")
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic)
        _para_format(p, line=line, before=before, after=after,
                     indent_cm=(1.0 if indent else None), align=align)
        return p

    def bullet(self, text, *, dash="−"):
        p = self.d.add_paragraph(style="Normal")
        run = p.add_run(f"{dash} {text}")
        _set_font(run)
        _para_format(p, indent_cm=1.0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        return p

    # -- headings ----------------------------------------------------------------
    def chapter_heading(self, title):
        """CHƯƠNG X: TÊN — in hoa, đậm, cỡ 14, căn giữa, sang trang mới."""
        self.chapter += 1
        self.fig = 0
        self.tbl = 0
        p = self.d.add_paragraph(style="Heading 1")
        p.paragraph_format.page_break_before = True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CHƯƠNG {self.chapter}: {title.upper()}")
        _set_font(run, size=SZ_CHAPTER, bold=True)
        return p

    def h2(self, num, title):
        p = self.d.add_paragraph(style="Heading 2")
        run = p.add_run(f"{num} {title}")
        _set_font(run, size=SZ_BODY, bold=True)
        _para_format(p, line=1.2, before=6, after=0)
        return p

    def h3(self, num, title):
        p = self.d.add_paragraph(style="Heading 3")
        run = p.add_run(f"{num} {title}")
        _set_font(run, size=SZ_BODY, bold=True)
        _para_format(p, line=1.2, before=6, after=0)
        return p

    def front_heading(self, title):
        """Tiêu đề phần đầu (LỜI CẢM ƠN…): in hoa, đậm 14, căn giữa, sang trang."""
        p = self.d.add_paragraph(style="Heading 1")
        p.paragraph_format.page_break_before = True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        _set_font(run, size=SZ_CHAPTER, bold=True)
        return p

    # -- figure / table caption --------------------------------------------------
    def figure(self, title):
        """Placeholder hình + caption 'Hình X.Y. ...' (dưới, căn giữa, cỡ 13)."""
        self.fig += 1
        ph = self.d.add_paragraph(style="Normal")
        r = ph.add_run("[Hình minh họa — sẽ chèn sau]")
        _set_font(r, italic=True)
        _para_format(ph, line=1.0, before=6, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        cap = self.d.add_paragraph(style="CaptionHinh")
        rc = cap.add_run(f"Hình {self.chapter}.{self.fig}. {title}")
        _set_font(rc, size=SZ_BODY, bold=False)
        _para_format(cap, line=1.0, before=2, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
        return cap

    def table_caption(self, title):
        """Caption 'Bảng X.Y. ...' (trên bảng, căn trái, cỡ 13)."""
        self.tbl += 1
        cap = self.d.add_paragraph(style="CaptionBang")
        rc = cap.add_run(f"Bảng {self.chapter}.{self.tbl}. {title}")
        _set_font(rc, size=SZ_BODY, bold=False)
        _para_format(cap, line=1.0, before=6, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
        return cap

    def data_table(self, headers, rows, widths=None):
        """Bảng dữ liệu (TNR 12, giãn dòng 1.0, khóa layout cố định, vừa khổ chữ)."""
        ncol = len(headers)
        # tổng độ rộng khả dụng = 21 - lề trái 3.5 - lề phải 2.0 = 15.5 cm
        if not widths:
            widths = [15.0 / ncol] * ncol
        total = sum(widths)
        if total > 15.4:                       # co lại cho vừa khổ chữ
            widths = [w * 15.4 / total for w in widths]
        t = self.d.add_table(rows=1, cols=ncol)
        t.style = "Table Grid"
        t.alignment = 1                         # căn giữa bảng
        t.autofit = False
        t.allow_autofit = False
        # khóa bố cục cố định để Word không tự co giãn cột
        tblPr = t._tbl.tblPr
        layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)

        def _fill(cell, text, *, bold):
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER   # chỉ canh giữa, không chỉnh paragraph
            r = p.add_run(text)
            _set_font(r, size=Pt(12), bold=bold)

        for i, h in enumerate(headers):
            _fill(t.rows[0].cells[i], h, bold=True)
        # lặp lại hàng tiêu đề khi bảng sang trang
        _repeat_header(t.rows[0])
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                _fill(cells[i], str(val), bold=False)
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
        return t

    def source_note(self, text):
        p = self.d.add_paragraph(style="Normal")
        r = p.add_run(text)
        _set_font(r, size=Pt(11), italic=True)
        _para_format(p, line=1.0, before=2, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
        return p


# ───────────────────────────── field helpers (TOC / TOF / PAGE) ────────────────
def _insert_field_after(ref_p_elem, instr, placeholder, document):
    """Chèn 1 paragraph chứa field (TOC/TOF) ngay sau ref_p_elem, trả về Paragraph."""
    from docx.text.paragraph import Paragraph
    newp = OxmlElement("w:p")
    ref_p_elem.addnext(newp)
    para = Paragraph(newp, document)
    r1 = para.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    r1._element.append(fb)
    r2 = para.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r2._element.append(it)
    r3 = para.add_run()
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    r3._element.append(fs)
    rp = para.add_run(placeholder)
    _set_font(rp, italic=True)
    r5 = para.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    r5._element.append(fe)
    _para_format(para, line=1.0, before=6, after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    return para


def _ensure_style(document, name, *, align, base="Normal"):
    """Tạo paragraph style cho caption nếu chưa có (để dựng Danh mục hình/bảng)."""
    from docx.enum.style import WD_STYLE_TYPE
    try:
        return document.styles[name]
    except KeyError:
        st = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = document.styles[base]
        st.font.name = FONT
        st.font.size = SZ_BODY
        st.font.bold = False
        st.paragraph_format.alignment = align
        return st


# ───────────────────────────── formatting compliance (309-QĐ) ──────────────────
def fix_global_formatting(document):
    """Lề + heading + đánh số trang theo 309-QĐ. KHÔNG đụng tới trang bìa (section 0)
    và KHÔNG đổi style Normal (vì trang bìa dùng Normal — đổi sẽ làm lệch bìa)."""
    # Heading sizes (H1=14, H2/H3=13 bold) — trang bìa không dùng heading nên an toàn
    for nm, sz in (("Heading 1", SZ_CHAPTER), ("Heading 2", SZ_BODY), ("Heading 3", SZ_BODY)):
        s = document.styles[nm]
        s.font.name = FONT
        s.font.size = sz
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
    # margins: CHỈ áp cho phần đầu (section 1) và thân bài (section 2) — GIỮ NGUYÊN bìa.
    for sec in document.sections[1:]:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.5)
        sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(1.0)
        sec.footer_distance = Cm(1.0)


def _set_pgnum(sec, fmt, start=None):
    sectpr = sec._sectPr
    pg = sectpr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectpr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def _footer_page_number(sec):
    """Đặt số trang căn giữa ở footer của section."""
    sec.footer.is_linked_to_previous = False
    f = sec.footer
    p = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_simple_page_field(p)


def _add_simple_page_field(paragraph):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._element.append(b); run._element.append(it); run._element.append(e)
    _set_font(run, size=SZ_BODY)


def setup_page_numbering(document):
    secs = document.sections
    # sec0 = title (no number), sec1 = front matter (roman), sec2 = body (arabic)
    if len(secs) >= 2:
        _set_pgnum(secs[1], "lowerRoman", start=1)
        _footer_page_number(secs[1])
    if len(secs) >= 3:
        _set_pgnum(secs[2], "decimal", start=1)
        _footer_page_number(secs[2])


# ───────────────────────────── body clearing + front matter ────────────────────
def _iter_body_children(document):
    return list(document.element.body)


def clear_body_after_frontmatter(document):
    """Xóa mọi <w:p>/<w:tbl> từ 'CHƯƠNG 1' tới hết (trừ sectPr cuối body)."""
    body = document.element.body
    # tìm paragraph 'CHƯƠNG 1' (Heading 1 chứa 'CHƯƠNG 1')
    from docx.text.paragraph import Paragraph
    start = None
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, document)
            if p.text.strip().upper().startswith("CHƯƠNG 1"):
                start = child
                break
    if start is None:
        raise RuntimeError("Không tìm thấy 'CHƯƠNG 1' trong template.")
    # xóa start và mọi anh em phía sau, NHƯNG giữ <w:sectPr> cuối (con trực tiếp body)
    to_remove = []
    node = start
    while node is not None:
        nxt = node.getnext()
        if node.tag == qn("w:sectPr"):
            break
        to_remove.append(node)
        node = nxt
    for n in to_remove:
        body.remove(n)


def find_para(document, predicate):
    from docx.text.paragraph import Paragraph
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, document)
            if predicate(p):
                return p
    return None


def rewrite_frontmatter(document):
    """Viết lại Lời cảm ơn + Lời cam kết; dựng Danh mục bảng/hình + Từ viết tắt."""
    from docx.text.paragraph import Paragraph

    title_full = ("Hệ thống tự động nhận dạng phong cách kiến trúc từ ảnh chụp "
                  "sử dụng trí tuệ nhân tạo")

    # --- LỜI CẢM ƠN: thay text các paragraph nội dung sau heading ---
    loi_camon = [
        "Lời đầu tiên, em xin bày tỏ lòng biết ơn sâu sắc đến thầy TS. Ngô Hồ Anh "
        "Khôi, giảng viên hướng dẫn khóa luận, đã tận tình định hướng, góp ý và "
        "đồng hành cùng em trong suốt quá trình thực hiện đề tài.",
        "Em xin chân thành cảm ơn quý thầy cô Khoa Công nghệ Thông tin, Trường Đại "
        "học Nam Cần Thơ đã truyền đạt cho em những kiến thức nền tảng và chuyên môn "
        "quý báu trong suốt thời gian học tập tại trường.",
        "Mặc dù đã nỗ lực trong quá trình nghiên cứu và thực hiện, do thời gian và "
        "kinh nghiệm còn hạn chế nên khóa luận không tránh khỏi những thiếu sót. Em "
        "rất mong nhận được sự góp ý của quý thầy cô để đề tài được hoàn thiện hơn.",
        "Em xin chân thành cảm ơn!",
    ]
    loi_camket = [
        f"Em xin cam kết khóa luận tốt nghiệp với đề tài “{title_full}” là công "
        "trình nghiên cứu của riêng em, được thực hiện dưới sự hướng dẫn của thầy "
        "TS. Ngô Hồ Anh Khôi.",
        "Các nội dung, kết quả nghiên cứu và sản phẩm được trình bày trong khóa luận "
        "là trung thực, do em tự thực hiện. Các tài liệu tham khảo được trích dẫn và "
        "ghi nguồn đầy đủ theo đúng quy định.",
        "Em xin hoàn toàn chịu trách nhiệm về tính trung thực và chính xác của các "
        "nội dung trong khóa luận này.",
    ]
    _replace_section_body(document, "LỜI CẢM ƠN", loi_camon)
    _replace_section_body(document, "LỜI CAM KẾT", loi_camket)


def _replace_section_body(document, heading_text, new_paras):
    """Thay các paragraph Normal nằm giữa heading_text và heading kế tiếp."""
    from docx.text.paragraph import Paragraph
    body = document.element.body
    children = [c for c in body.iterchildren() if c.tag == qn("w:p")]
    # locate heading
    idx = None
    plist = [Paragraph(c, document) for c in children]
    for i, p in enumerate(plist):
        if p.text.strip().upper() == heading_text.upper() and p.style.name.startswith("Heading"):
            idx = i
            break
    if idx is None:
        return
    # collect following Normal paragraphs until next Heading
    j = idx + 1
    body_paras = []
    while j < len(plist):
        st = plist[j].style.name
        if st.startswith("Heading"):
            break
        body_paras.append(plist[j])
        j += 1
    # rewrite: reuse existing paragraphs where possible, else insert
    ref = plist[idx]._p
    # remove old body paragraphs (but keep any carrying sectPr)
    for p in body_paras:
        if p._p.find(qn("w:pPr")) is not None and p._p.find(qn("w:pPr")).find(qn("w:sectPr")) is not None:
            continue
        p._p.getparent().remove(p._p)
    # insert new paragraphs after heading
    cursor = ref
    for txt in new_paras:
        newp = OxmlElement("w:p")
        cursor.addnext(newp)
        cursor = newp
        para = Paragraph(newp, document)
        run = para.add_run(txt)
        _set_font(run)
        _para_format(para, indent_cm=1.0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _style_front_title(document, text):
    """Chuyển tiêu đề phần đầu (DANH MỤC…) sang Heading 1 căn giữa, 14pt."""
    from docx.text.paragraph import Paragraph
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, document)
            if p.text.strip().upper() == text.upper():
                p.style = document.styles["Heading 1"]
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.page_break_before = True
                for r in p.runs:
                    _set_font(r, size=SZ_CHAPTER, bold=True)
                return p
    return None


ABBREVIATIONS = [
    ("AI", "Trí tuệ nhân tạo (Artificial Intelligence)"),
    ("API", "Giao diện lập trình ứng dụng (Application Programming Interface)"),
    ("CNN", "Mạng nơ-ron tích chập (Convolutional Neural Network)"),
    ("CoT", "Chuỗi suy luận (Chain-of-Thought)"),
    ("CSDL", "Cơ sở dữ liệu"),
    ("ERD", "Mô hình thực thể – quan hệ (Entity–Relationship Diagram)"),
    ("JSON", "Định dạng trao đổi dữ liệu (JavaScript Object Notation)"),
    ("JWT", "Mã thông báo web JSON (JSON Web Token)"),
    ("KB", "Cơ sở tri thức (Knowledge Base)"),
    ("LLM", "Mô hình ngôn ngữ lớn (Large Language Model)"),
    ("OAuth", "Giao thức ủy quyền mở (Open Authorization)"),
    ("ORM", "Ánh xạ đối tượng – quan hệ (Object–Relational Mapping)"),
    ("RBAC", "Phân quyền theo vai trò (Role-Based Access Control)"),
    ("REST", "Kiến trúc chuyển trạng thái biểu diễn (Representational State Transfer)"),
    ("UML", "Ngôn ngữ mô hình hóa thống nhất (Unified Modeling Language)"),
    ("VLM", "Mô hình Thị giác – Ngôn ngữ (Vision–Language Model)"),
]


def build_frontmatter_tail(document):
    """Chèn trường TOC/TOF + danh sách viết tắt và di chuyển section break."""
    from docx.text.paragraph import Paragraph

    # chuẩn hóa tiêu đề phần đầu
    for tt in ("DANH MỤC BẢNG", "DANH MỤC HÌNH", "DANH MỤC TỪ VIẾT TẮT"):
        _style_front_title(document, tt)

    def find_title(text):
        for child in document.element.body.iterchildren():
            if child.tag == qn("w:p"):
                p = Paragraph(child, document)
                if p.text.strip().upper() == text.upper():
                    return p
        return None

    # mỗi mục phần đầu sang trang mới (trừ LỜI CẢM ƠN — đã ở đầu section sau bìa)
    for tt in ("LỜI CAM KẾT", "MỤC LỤC"):
        ph = find_title(tt)
        if ph is not None:
            ph.paragraph_format.page_break_before = True

    # MỤC LỤC → TOC field
    ml = find_title("MỤC LỤC")
    if ml is not None:
        _insert_field_after(ml._p, 'TOC \\o "1-3" \\h \\z \\u ',
                            "Mục lục tự động — nhấn Ctrl+A rồi F9 để cập nhật.", document)

    # DANH MỤC BẢNG → Table of Figures dựng theo style CaptionBang
    db = find_title("DANH MỤC BẢNG")
    if db is not None:
        _insert_field_after(db._p, 'TOC \\h \\z \\t "CaptionBang,1" ',
                            "Danh mục bảng tự động — nhấn Ctrl+A rồi F9 để cập nhật.", document)

    dh = find_title("DANH MỤC HÌNH")
    if dh is not None:
        _insert_field_after(dh._p, 'TOC \\h \\z \\t "CaptionHinh,1" ',
                            "Danh mục hình tự động — nhấn Ctrl+A rồi F9 để cập nhật.", document)

    # DANH MỤC TỪ VIẾT TẮT → bảng viết tắt + chuyển sectPr
    dv = find_title("DANH MỤC TỪ VIẾT TẮT")
    if dv is not None:
        cursor = dv._p
        last = None
        for abbr, mean in ABBREVIATIONS:
            newp = OxmlElement("w:p")
            cursor.addnext(newp)
            cursor = newp
            para = Paragraph(newp, document)
            r1 = para.add_run(f"{abbr}")
            _set_font(r1, bold=True)
            r2 = para.add_run(f"\t: {mean}")
            _set_font(r2)
            _para_format(para, line=1.2, before=2, after=2,
                         indent_cm=None, align=WD_ALIGN_PARAGRAPH.LEFT)
            last = newp
        if last is not None:
            _move_sectpr(dv._p, last)


def build_chapters(doc):
    chuong_1(doc)
    chuong_2(doc)
    chuong_3(doc)
    chuong_4(doc)
    chuong_5(doc)
    chuong_6(doc)
    tai_lieu_tham_khao(doc)


# ════════════════════════════ CHƯƠNG 1 — GIỚI THIỆU ════════════════════════════
def chuong_1(d):
    d.chapter_heading("Giới thiệu")

    d.h2("1.1.", "Đặt vấn đề và lý do chọn đề tài")
    d.para("Phong cách kiến trúc là sự kết tinh của lịch sử, văn hóa và kỹ thuật xây "
           "dựng của một thời kỳ và một vùng đất. Việc nhận biết phong cách của một công "
           "trình không chỉ có ý nghĩa học thuật mà còn phục vụ thiết thực cho du lịch, "
           "giáo dục, bảo tồn di sản và các ứng dụng tra cứu thông tin. Tuy nhiên, để "
           "phân biệt chính xác hàng trăm phong cách kiến trúc trên thế giới, người xem "
           "thường phải là chuyên gia được đào tạo bài bản, trong khi phần lớn người dùng "
           "phổ thông chỉ có thể nhận ra một vài phong cách quen thuộc.")
    d.para("Trong lĩnh vực thị giác máy tính, bài toán nhận dạng phong cách kiến trúc đã "
           "được nghiên cứu từ sớm, điển hình là công trình của Xu và cộng sự (2014) phân "
           "loại 25 phong cách bằng mô hình học máy. Tuy vậy, hướng tiếp cận phân loại "
           "đóng (closed-set) — huấn luyện mô hình trên một số lớp cố định — gặp hạn chế "
           "căn bản: độ chính xác suy giảm nhanh khi số lớp tăng lên, và muốn bổ sung một "
           "phong cách mới thì phải thu thập dữ liệu rồi huấn luyện lại toàn bộ. Với hơn "
           "một trăm phong cách kiến trúc trải dài từ cổ đại đến đương đại, nhiều phong "
           "cách lại rất hiếm ảnh có nhãn, cách làm này gần như bất khả thi.")
    d.para("Trong những năm gần đây, sự phát triển của các mô hình ngôn ngữ lớn (Large "
           "Language Model — LLM) và đặc biệt là các mô hình thị giác – ngôn ngữ "
           "(Vision–Language Model — VLM) như GPT-4o, Gemini hay Grok đã mở ra khả năng "
           "phân tích hình ảnh kèm suy luận bằng ngôn ngữ tự nhiên mà không cần huấn "
           "luyện riêng cho từng lớp. Đây là tiền đề để xây dựng một hệ thống nhận dạng "
           "phong cách kiến trúc theo hướng mở từ vựng (open-vocabulary): có thể nhận "
           "biết một tập phong cách rất rộng và dễ dàng mở rộng bằng cách bổ sung tri "
           "thức thay vì huấn luyện lại.")
    d.para("Tuy nhiên, bản thân một mô hình ngôn ngữ – thị giác đơn lẻ vẫn dễ mắc lỗi "
           "“ảo giác” (hallucination), tự tin sai và thiếu nhất quán giữa các lần chạy. "
           "Xuất phát từ thực trạng đó, đề tài “Hệ thống tự động nhận dạng phong cách "
           "kiến trúc từ ảnh chụp sử dụng trí tuệ nhân tạo” được thực hiện nhằm xây dựng "
           "một hệ thống nhận dạng phong cách kiến trúc theo hướng mở từ vựng, kết hợp "
           "nhiều mô hình theo cơ chế hội đồng – trọng tài, có nối khớp với cơ sở tri "
           "thức và có khả năng từ chối trả lời khi không đủ chắc chắn.")

    d.h2("1.2.", "Mục tiêu đề tài")
    d.h3("1.2.1.", "Mục tiêu tổng quát")
    d.para("Mục tiêu tổng quát của đề tài là xây dựng một hệ thống phần mềm hoàn chỉnh, "
           "cho phép người dùng tải lên ảnh chụp một công trình kiến trúc và nhận về phong "
           "cách kiến trúc của công trình kèm theo lời giải thích dựa trên bằng chứng "
           "quan sát được, theo hướng nhận dạng mở từ vựng và có độ tin cậy cao.")
    d.h3("1.2.2.", "Mục tiêu cụ thể")
    d.para("Để đạt được mục tiêu tổng quát, đề tài xác định các mục tiêu cụ thể sau:")
    d.bullet("Xây dựng cơ sở tri thức về phong cách kiến trúc gồm 106 phong cách thuộc 12 "
             "họ, có mô tả đặc trưng và nguồn tham chiếu, làm nền tảng cho việc nối khớp "
             "và giải thích.")
    d.bullet("Thiết kế và hiện thực một quy trình phân tích đa tác tử (multi-agent): trích "
             "xuất bằng chứng từ ảnh, sinh tập ứng viên từ cơ sở tri thức, chấm điểm bằng "
             "hội đồng nhiều mô hình thị giác và tổng hợp kết quả bằng một trọng tài.")
    d.bullet("Tích hợp cơ chế từ chối trả lời (abstention) dựa trên mức đồng thuận giữa các "
             "giám khảo, giúp hệ thống biết “nói không chắc” thay vì đoán bừa.")
    d.bullet("Xây dựng ứng dụng web hoàn chỉnh: xác thực bằng Google, trang phân tích và "
             "lịch sử cho người dùng, trang quản trị cho quản trị viên, hỗ trợ song ngữ "
             "Việt – Anh.")
    d.bullet("Đánh giá định lượng hệ thống trên bộ dữ liệu thử nghiệm 200 ảnh và đối sánh "
             "với hai mô hình nền trên cùng một thang đo.")

    d.h2("1.3.", "Đối tượng và phạm vi đề tài")
    d.h3("1.3.1.", "Đối tượng nghiên cứu")
    d.para("Đối tượng nghiên cứu của đề tài là phương pháp nhận dạng phong cách kiến trúc "
           "từ ảnh chụp theo hướng mở từ vựng, dựa trên sự phối hợp giữa nhiều mô hình thị "
           "giác – ngôn ngữ và một cơ sở tri thức phong cách. Đối tượng sử dụng trực tiếp "
           "của hệ thống gồm hai nhóm: người dùng phổ thông (tải ảnh lên để tra cứu phong "
           "cách) và quản trị viên (quản lý người dùng, dữ liệu và theo dõi hoạt động).")
    d.h3("1.3.2.", "Phạm vi đề tài")
    d.para("Về phạm vi chức năng, đề tài tập trung vào việc nhận dạng phong cách kiến trúc "
           "ở mức tổng thể công trình, kèm giải thích và lịch sử phân tích; không nhằm mục "
           "tiêu định vị chi tiết từng cấu kiện hay đo đạc kích thước công trình.")
    d.para("Về phạm vi công nghệ, hệ thống được xây dựng dưới dạng ứng dụng web với phần "
           "thân (backend) viết bằng Python và FastAPI, phần giao diện (frontend) viết "
           "bằng React, cơ sở dữ liệu SQL Server, và sử dụng dịch vụ của các nhà cung cấp "
           "mô hình ngôn ngữ – thị giác qua giao diện lập trình ứng dụng (API).")
    d.para("Về phạm vi dữ liệu thử nghiệm, đề tài đánh giá trên bộ 200 ảnh tương ứng 100 "
           "phong cách (mỗi phong cách 2 ảnh được sinh từ hai nguồn khác nhau), có nhãn "
           "phong cách chính làm đáp án tham chiếu.")

    d.h2("1.4.", "Phương pháp thực hiện")
    d.para("Đề tài được thực hiện theo quy trình phát triển phần mềm kết hợp nghiên cứu "
           "thực nghiệm, gồm các bước chính sau:")
    d.bullet("Khảo sát và phân tích yêu cầu: xác định mục tiêu, đối tượng sử dụng, phạm vi "
             "và các yêu cầu chức năng, phi chức năng của hệ thống.")
    d.bullet("Nghiên cứu cơ sở lý thuyết: tìm hiểu các mô hình thị giác – ngôn ngữ, mô hình "
             "ngôn ngữ lớn, hệ đa tác tử, cơ sở tri thức và các công trình liên quan.")
    d.bullet("Thiết kế hệ thống: thiết kế cơ sở dữ liệu bằng mô hình thực thể – quan hệ, "
             "thiết kế kiến trúc và các sơ đồ UML, thiết kế quy trình phân tích đa tác tử.")
    d.bullet("Hiện thực hệ thống: xây dựng cơ sở tri thức, lập trình quy trình phân tích, "
             "phần thân, phần giao diện và cơ chế xác thực – phân quyền.")
    d.bullet("Thử nghiệm và đánh giá: xây dựng quy trình chấm điểm tự động, đo các chỉ số "
             "định lượng trên bộ dữ liệu 200 ảnh và đối sánh với hai mô hình nền.")

    d.h2("1.5.", "Ý nghĩa khoa học và thực tiễn")
    d.h3("1.5.1.", "Ý nghĩa khoa học")
    d.para("Về mặt khoa học, đề tài đề xuất một kiến trúc nhận dạng phong cách kiến trúc "
           "mở từ vựng, trong đó nhiều mô hình thị giác – ngôn ngữ đóng vai trò hội đồng "
           "giám khảo độc lập, được nối khớp với cơ sở tri thức và có cơ chế từ chối dựa "
           "trên mức đồng thuận. Đây là một minh chứng thực nghiệm cho hướng kết hợp nhiều "
           "mô hình ngôn ngữ lớn (hội đồng – trọng tài) áp dụng vào bài toán nhận dạng "
           "ảnh chi tiết (fine-grained), thay cho cách huấn luyện phân loại đóng truyền thống.")
    d.h3("1.5.2.", "Ý nghĩa thực tiễn")
    d.para("Về mặt thực tiễn, hệ thống cung cấp một công cụ trực quan giúp người dùng phổ "
           "thông tra cứu nhanh phong cách của một công trình kiến trúc kèm giải thích dễ "
           "hiểu bằng tiếng Việt, hỗ trợ học tập, du lịch và tìm hiểu di sản. Khả năng mở "
           "rộng phong cách bằng cách bổ sung tri thức (không cần huấn luyện lại) giúp hệ "
           "thống dễ bảo trì và phát triển lâu dài.")

    d.h2("1.6.", "Đóng góp của đề tài")
    d.para("Đề tài mang lại các đóng góp cụ thể sau:")
    d.bullet("Một cơ sở tri thức phong cách kiến trúc gồm 106 phong cách thuộc 12 họ, có "
             "đặc trưng nhận dạng và nguồn tham chiếu.")
    d.bullet("Một quy trình phân tích đa tác tử mở từ vựng kết hợp trích xuất bằng chứng, "
             "nối khớp tri thức, hội đồng giám khảo thị giác và trọng tài.")
    d.bullet("Một cơ chế từ chối trả lời và báo cáo độ tin cậy dựa trên mức đồng thuận, "
             "giúp hạn chế kết luận sai một cách tự tin.")
    d.bullet("Một ứng dụng web hoàn chỉnh, đa vai trò, song ngữ, cùng bộ kết quả đánh giá "
             "định lượng trên 200 ảnh có đối sánh với hai mô hình nền.")

    d.h2("1.7.", "Cấu trúc khóa luận")
    d.para("Ngoài phần Giới thiệu, nội dung khóa luận được tổ chức thành sáu chương:")
    d.bullet("Chương 1 — Giới thiệu: trình bày lý do chọn đề tài, mục tiêu, đối tượng, phạm "
             "vi, phương pháp thực hiện, ý nghĩa và đóng góp.")
    d.bullet("Chương 2 — Cơ sở lý thuyết: trình bày các kiến thức nền tảng về phong cách "
             "kiến trúc, mô hình thị giác – ngôn ngữ, hệ đa tác tử, cơ sở tri thức và công "
             "nghệ sử dụng.")
    d.bullet("Chương 3 — Thiết kế cơ sở dữ liệu: trình bày mô hình thực thể – quan hệ và mô "
             "tả chi tiết các bảng dữ liệu của hệ thống.")
    d.bullet("Chương 4 — Phân tích và thiết kế hệ thống: trình bày yêu cầu, kiến trúc tổng "
             "thể, quy trình phân tích đa tác tử và các sơ đồ UML.")
    d.bullet("Chương 5 — Môi trường và kết quả thử nghiệm: trình bày môi trường triển khai, "
             "phương pháp đánh giá, kết quả định lượng và giao diện hệ thống.")
    d.bullet("Chương 6 — Kết quả và hướng phát triển của đề tài: tổng kết kết quả đạt được, "
             "hạn chế và hướng phát triển.")


# ════════════════════════════ CHƯƠNG 2 — CƠ SỞ LÝ THUYẾT ═══════════════════════
def chuong_2(d):
    d.chapter_heading("Cơ sở lý thuyết")

    d.h2("2.1.", "Giới thiệu chung")
    d.para("Chương này trình bày các cơ sở lý thuyết được sử dụng làm nền tảng cho việc "
           "phân tích, thiết kế và hiện thực hệ thống. Nội dung gồm bốn nhóm: kiến thức "
           "về phong cách kiến trúc và bài toán nhận dạng; các mô hình trí tuệ nhân tạo "
           "phục vụ phân tích ảnh và suy luận (mô hình thị giác – ngôn ngữ, mô hình ngôn "
           "ngữ lớn, hệ đa tác tử, cơ sở tri thức); các công nghệ phần mềm và cơ chế xác "
           "thực; và các mô hình dùng trong phân tích – thiết kế hệ thống.")

    d.h2("2.2.", "Phong cách kiến trúc và bài toán nhận dạng")
    d.para("Phong cách kiến trúc (architectural style) là tập hợp các đặc trưng hình thức "
           "đặc trưng cho một thời kỳ, một vùng hoặc một trào lưu, thể hiện qua khối hình "
           "tổng thể, mái, hệ cột, dạng vòm cửa, cách trang trí mặt đứng và vật liệu. Các "
           "phong cách có quan hệ phả hệ với nhau và thường được tổ chức thành các họ "
           "(family) như cổ đại, trung cổ châu Âu, Hồi giáo, Đông Á… Việc phân biệt các "
           "phong cách gần nhau (ví dụ Romanesque và Romanesque Revival) là một bài toán "
           "nhận dạng chi tiết (fine-grained) khó, ngay cả với chuyên gia.")
    d.para("Bài toán nhận dạng phong cách có thể tiếp cận theo hai hướng. Hướng phân loại "
           "đóng (closed-set) huấn luyện một mô hình trên một số lớp cố định; ưu điểm là "
           "đơn giản nhưng nhược điểm là không mở rộng được và độ chính xác giảm khi số "
           "lớp tăng. Hướng mở từ vựng (open-vocabulary) cho phép nhận biết một tập phong "
           "cách rất rộng dựa trên mô tả bằng ngôn ngữ và tri thức, dễ mở rộng bằng cách "
           "bổ sung tri thức thay vì huấn luyện lại. Đề tài này theo hướng mở từ vựng.")
    d.para("Để bảo đảm tính chuẩn xác và truy vết được nguồn gốc, danh mục phong cách và "
           "đặc trưng trong hệ thống được tham chiếu đến các nguồn uy tín như Bộ từ chuẩn "
           "Nghệ thuật và Kiến trúc của Getty (Art & Architecture Thesaurus — AAT) và cơ "
           "sở dữ liệu tri thức mở Wikidata, cùng các tài liệu lịch sử kiến trúc.")

    d.h2("2.3.", "Thị giác máy tính và mô hình thị giác – ngôn ngữ")
    d.para("Thị giác máy tính (computer vision) là lĩnh vực giúp máy tính “hiểu” nội dung "
           "hình ảnh. Trong học sâu, mạng nơ-ron tích chập (Convolutional Neural Network — "
           "CNN), tiêu biểu là kiến trúc ResNet (He và cộng sự, 2016), là nền tảng cho "
           "nhiều bài toán phân loại ảnh. Kiến trúc Transformer (Vaswani và cộng sự, 2017) "
           "và biến thể cho ảnh là Vision Transformer (ViT) về sau đã trở thành xương sống "
           "của nhiều mô hình hiện đại.")
    d.para("Mô hình thị giác – ngôn ngữ (Vision–Language Model — VLM) là lớp mô hình có khả "
           "năng tiếp nhận đồng thời hình ảnh và văn bản, từ đó mô tả, trả lời câu hỏi hay "
           "suy luận về nội dung ảnh bằng ngôn ngữ tự nhiên. Mô hình CLIP (Radford và cộng "
           "sự, 2021) là một dấu mốc khi học liên kết ảnh – văn bản ở quy mô lớn, cho phép "
           "phân loại không cần huấn luyện lại (zero-shot). Các mô hình thế hệ mới như "
           "GPT-4o (OpenAI), Gemini (Google) và Grok (xAI) có năng lực thị giác mạnh, có "
           "thể nhận một ảnh công trình và trực tiếp gọi tên phong cách kèm lý giải. Hệ "
           "thống trong đề tài sử dụng đồng thời ba mô hình thị giác – ngôn ngữ này như ba "
           "giám khảo độc lập.")

    d.h2("2.4.", "Mô hình ngôn ngữ lớn và hệ đa tác tử")
    d.para("Mô hình ngôn ngữ lớn (Large Language Model — LLM) là mô hình học sâu được huấn "
           "luyện trên khối văn bản khổng lồ, có khả năng sinh và suy luận bằng ngôn ngữ "
           "tự nhiên. Một kỹ thuật quan trọng là chuỗi suy luận (Chain-of-Thought — CoT), "
           "trong đó mô hình được yêu cầu trình bày các bước lập luận trước khi đưa ra kết "
           "luận, giúp tăng tính nhất quán giữa lập luận và quyết định.")
    d.para("Hạn chế cố hữu của một mô hình đơn lẻ là hiện tượng “ảo giác” (hallucination — "
           "đưa ra thông tin nghe hợp lý nhưng sai) và sự thiếu ổn định giữa các lần chạy. "
           "Một hướng khắc phục là hệ đa tác tử (multi-agent system): nhiều mô hình cùng "
           "tham gia, kiểm tra chéo và tổng hợp kết quả. Du và cộng sự (2024) cho thấy cơ "
           "chế tranh luận đa tác tử (multi-agent debate) giúp tăng tính xác thực và năng "
           "lực suy luận; Wang và cộng sự (2024) đề xuất phương pháp Hỗn hợp Tác tử "
           "(Mixture-of-Agents) kết hợp thế mạnh của nhiều mô hình. Kế thừa ý tưởng này, "
           "đề tài tổ chức nhiều mô hình thành một hội đồng giám khảo độc lập, đo mức đồng "
           "thuận giữa chúng và dùng một mô hình mạnh làm trọng tài để đưa ra kết luận cuối.")

    d.h2("2.5.", "Cơ sở tri thức và cơ chế nối khớp")
    d.para("Cơ sở tri thức (Knowledge Base — KB) là một kho thông tin có cấu trúc về các "
           "phong cách: tên chuẩn, tên đồng nghĩa (alias), họ, vùng – thời kỳ, đặc trưng "
           "nhận dạng và nguồn tham chiếu. Cơ sở tri thức giúp “neo” kết quả của mô hình "
           "vào tri thức có kiểm chứng (grounding), tránh việc mô hình tự bịa ra phong "
           "cách không tồn tại.")
    d.para("Cơ chế nối khớp (matching) ánh xạ một tên phong cách do mô hình đề xuất về một "
           "mục trong cơ sở tri thức theo ba mức: khớp chính xác theo tên, khớp theo tên "
           "đồng nghĩa, và khớp xấp xỉ (fuzzy matching) dựa trên độ tương đồng chuỗi khi "
           "hai mức trên không thành công. Nhờ tên đồng nghĩa, các cách gọi khác nhau của "
           "cùng một phong cách (ví dụ “Chinese” và “Traditional Chinese”) được gộp về "
           "cùng một mục. Khi mức đồng thuận giữa các giám khảo thấp, hệ thống áp dụng cơ "
           "chế từ chối (abstention) — báo “không chắc chắn” thay vì đưa ra kết luận rủi ro.")

    d.h2("2.6.", "Kiến trúc phần mềm và công nghệ sử dụng")
    d.para("Hệ thống được xây dựng theo kiến trúc nhiều tầng, tách biệt phần giao diện, "
           "phần thân xử lý nghiệp vụ và tầng dữ liệu. Các công nghệ chính gồm:")
    d.bullet("FastAPI: một khung phát triển web bằng Python theo phong cách bất đồng bộ "
             "(asynchronous), dùng để xây dựng các giao diện lập trình ứng dụng (API) theo "
             "kiến trúc REST (Representational State Transfer).")
    d.bullet("SQLAlchemy (bất đồng bộ) và trình điều khiển aioodbc: lớp ánh xạ đối tượng – "
             "quan hệ (Object–Relational Mapping — ORM) giúp thao tác cơ sở dữ liệu SQL "
             "Server bằng đối tượng Python.")
    d.bullet("React và Vite: thư viện và công cụ xây dựng giao diện người dùng phía trình "
             "duyệt, hỗ trợ giao diện song ngữ và chế độ sáng/tối.")
    d.bullet("SQL Server: hệ quản trị cơ sở dữ liệu quan hệ dùng để lưu trữ người dùng, dự "
             "án, ảnh, kết quả phân tích và nhật ký hệ thống.")
    d.bullet("Dịch vụ mô hình ngôn ngữ – thị giác (Gemini, OpenAI GPT-4o, Grok, DeepSeek) "
             "được gọi qua API, đảm nhiệm vai trò trích xuất bằng chứng, giám khảo, trọng "
             "tài và dịch thuật.")

    d.h2("2.7.", "Cơ chế xác thực và phân quyền")
    d.para("Xác thực (authentication) là quá trình xác minh danh tính người dùng, còn phân "
           "quyền (authorization) quyết định người dùng được phép làm gì. Hệ thống sử dụng "
           "đăng nhập một chạm bằng tài khoản Google theo chuẩn OAuth 2.0 (Open "
           "Authorization): phần thân hệ thống xác minh mã thông báo do Google cấp rồi "
           "phát hành mã thông báo web JSON (JSON Web Token — JWT) ký bằng thuật toán "
           "HS256. Trình duyệt người dùng không bao giờ lưu trực tiếp mã của Google.")
    d.para("Mỗi JWT chứa thông tin định danh và vai trò người dùng. Việc phân quyền theo "
           "vai trò (Role-Based Access Control — RBAC) phân biệt hai vai trò chính là "
           "người dùng và quản trị viên; các chức năng quản trị chỉ truy cập được khi mã "
           "thông báo mang vai trò quản trị viên và tài khoản còn hiệu lực.")

    d.h2("2.8.", "Các mô hình dùng trong phân tích và thiết kế hệ thống")
    d.para("Để phân tích và thiết kế hệ thống, đề tài sử dụng Ngôn ngữ mô hình hóa thống "
           "nhất (Unified Modeling Language — UML) và mô hình thực thể – quan hệ:")
    d.bullet("Sơ đồ ca sử dụng (Use Case Diagram): mô tả chức năng hệ thống từ góc nhìn của "
             "các tác nhân (người dùng, quản trị viên).")
    d.bullet("Sơ đồ hoạt động (Activity Diagram): mô tả luồng xử lý của một quy trình "
             "nghiệp vụ như đăng nhập hay phân tích ảnh.")
    d.bullet("Sơ đồ trình tự (Sequence Diagram): mô tả thứ tự tương tác giữa các thành phần "
             "theo thời gian.")
    d.bullet("Sơ đồ lớp, gói và thành phần (Class, Package, Component Diagram): mô tả cấu "
             "trúc tĩnh và cách tổ chức các thành phần phần mềm.")
    d.bullet("Sơ đồ triển khai (Deployment Diagram): mô tả cách phân bố các thành phần trên "
             "hạ tầng vận hành.")
    d.bullet("Mô hình thực thể – quan hệ (Entity–Relationship Diagram — ERD): mô tả các "
             "thực thể dữ liệu và quan hệ giữa chúng, làm cơ sở cho thiết kế cơ sở dữ liệu.")

    d.h2("2.9.", "Kết chương")
    d.para("Chương 2 đã trình bày các cơ sở lý thuyết nền tảng: đặc điểm của bài toán nhận "
           "dạng phong cách kiến trúc và lợi thế của hướng mở từ vựng; các mô hình thị "
           "giác – ngôn ngữ và mô hình ngôn ngữ lớn; ý tưởng hệ đa tác tử và cơ sở tri "
           "thức; cùng các công nghệ và mô hình thiết kế được sử dụng. Trên cơ sở đó, các "
           "chương tiếp theo trình bày thiết kế cơ sở dữ liệu (Chương 3) và phân tích – "
           "thiết kế hệ thống (Chương 4).")


# ════════════════════════════ CHƯƠNG 3 — THIẾT KẾ CƠ SỞ DỮ LIỆU ════════════════
_DD_HEADERS = ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"]


def _dd(d, caption, rows):
    d.table_caption(caption)
    d.data_table(_DD_HEADERS, rows, widths=[3.6, 3.4, 3.0, 6.0])


def chuong_3(d):
    d.chapter_heading("Thiết kế cơ sở dữ liệu")

    d.h2("3.1.", "Mô hình thực thể – quan hệ tổng thể")
    d.para("Cơ sở dữ liệu của hệ thống được xây dựng trên SQL Server, gồm 14 bảng tổ chức "
           "thành năm nhóm: nhóm người dùng và phân quyền; nhóm ảnh và thành phần; nhóm "
           "tác tử và lần chạy; nhóm kết quả phân tích; và nhóm nhật ký hệ thống. Các bảng "
           "liên kết với nhau qua khóa ngoại, trong đó toàn bộ khóa chính sử dụng kiểu định "
           "danh duy nhất (UNIQUEIDENTIFIER) do hệ quản trị tự sinh. Quan hệ tổng thể được "
           "thể hiện qua mô hình thực thể – quan hệ trong Hình 3.1.")
    d.figure("Mô hình thực thể – quan hệ (ERD) tổng thể của hệ thống")
    d.para("Luồng quan hệ chính: mỗi vai trò (Roles) gắn với nhiều người dùng (Users); mỗi "
           "người dùng sở hữu nhiều dự án (Projects); mỗi dự án chứa nhiều ảnh (Images); "
           "mỗi ảnh có thể sinh ra nhiều thành phần (Components), các lần chạy tác tử "
           "(AgentRuns), các giả thuyết phong cách (BuildingStyleHypotheses) và một kết "
           "quả phong cách cuối cùng (BuildingStyleResults).")

    d.h2("3.2.", "Mô tả chi tiết các bảng dữ liệu")

    d.h3("3.2.1.", "Nhóm bảng Người dùng và Phân quyền")
    _dd(d, "Cấu trúc bảng Roles (vai trò người dùng)", [
        ["RoleId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh vai trò, hệ tự sinh."],
        ["RoleName", "NVARCHAR(50)", "NOT NULL", "Tên vai trò: 'user' hoặc 'admin'."],
    ])
    _dd(d, "Cấu trúc bảng Users (người dùng)", [
        ["UserId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh người dùng."],
        ["Email", "NVARCHAR(255)", "NOT NULL, UNIQUE", "Địa chỉ email đăng nhập."],
        ["Name", "NVARCHAR(255)", "NOT NULL", "Họ tên hiển thị (có dấu)."],
        ["Picture", "NVARCHAR(500)", "NULL", "Đường dẫn ảnh đại diện."],
        ["GoogleSub", "NVARCHAR(255)", "NULL, UNIQUE", "Định danh tài khoản Google."],
        ["IsActive", "BIT", "NOT NULL, mặc định 1", "Trạng thái hoạt động của tài khoản."],
        ["RoleId", "UNIQUEIDENTIFIER", "Khóa ngoại → Roles", "Vai trò của người dùng."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
        ["UpdatedAt", "DATETIME", "NULL", "Thời điểm cập nhật gần nhất."],
        ["PasswordHash", "NVARCHAR(255)", "NULL", "Mã băm mật khẩu (nếu đăng nhập nội bộ)."],
    ])
    _dd(d, "Cấu trúc bảng Projects (dự án)", [
        ["ProjectId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh dự án."],
        ["UserId", "UNIQUEIDENTIFIER", "Khóa ngoại → Users", "Chủ sở hữu dự án."],
        ["ProjectName", "NVARCHAR(255)", "NULL", "Tên dự án (có dấu)."],
        ["Description", "NVARCHAR(MAX)", "NULL", "Mô tả dự án."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
    ])

    d.h3("3.2.2.", "Nhóm bảng Ảnh và Thành phần")
    _dd(d, "Cấu trúc bảng Images (ảnh phân tích)", [
        ["ImageId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh ảnh."],
        ["ProjectId", "UNIQUEIDENTIFIER", "Khóa ngoại → Projects", "Dự án chứa ảnh."],
        ["ImagePath", "NVARCHAR(500)", "NULL", "Đường dẫn lưu ảnh (đặt tên theo UUID)."],
        ["AnalysisStatus", "NVARCHAR(50)", "NOT NULL, mặc định 'pending'", "Trạng thái: pending | processing | completed | failed."],
        ["ErrorMessage", "NVARCHAR(MAX)", "NULL", "Thông báo lỗi nếu phân tích thất bại."],
        ["UpdatedAt", "DATETIME", "NULL", "Thời điểm cập nhật gần nhất."],
        ["UploadedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tải lên."],
    ])
    _dd(d, "Cấu trúc bảng Components (thành phần kiến trúc)", [
        ["ComponentId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh thành phần."],
        ["ImageId", "UNIQUEIDENTIFIER", "Khóa ngoại → Images", "Ảnh chứa thành phần."],
        ["ComponentType", "NVARCHAR(100)", "NULL", "Loại thành phần kiến trúc."],
        ["BoundingBox", "NVARCHAR(255)", "NULL", "Tọa độ khung bao của thành phần (JSON)."],
        ["CropImageUrl", "NVARCHAR(500)", "NULL", "Đường dẫn ảnh cắt của thành phần."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
    ])

    d.h3("3.2.3.", "Nhóm bảng Tác tử và Lần chạy")
    _dd(d, "Cấu trúc bảng Agents (tác tử trong quy trình)", [
        ["AgentId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh tác tử."],
        ["AgentName", "NVARCHAR(100)", "NULL", "Tên tác tử (ví dụ: Arbiter, Panel: Gemini)."],
        ["Description", "NVARCHAR(500)", "NULL", "Mô tả vai trò và mô hình sử dụng."],
    ])
    _dd(d, "Cấu trúc bảng AgentRuns (lần chạy tác tử)", [
        ["RunId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh lần chạy."],
        ["ImageId", "UNIQUEIDENTIFIER", "Khóa ngoại → Images, NULL", "Ảnh tương ứng (mức công trình)."],
        ["ComponentId", "UNIQUEIDENTIFIER", "Khóa ngoại → Components, NULL", "Thành phần tương ứng (nếu có)."],
        ["AgentId", "UNIQUEIDENTIFIER", "Khóa ngoại → Agents", "Tác tử thực thi."],
        ["InputData", "NVARCHAR(MAX)", "NULL", "Dữ liệu/đầu vào đưa cho tác tử."],
        ["OutputData", "NVARCHAR(MAX)", "NULL", "Kết quả tổng hợp của tác tử."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
        ["RawOutput", "NVARCHAR(MAX)", "NULL", "Phản hồi thô của mô hình trước khi phân tích."],
        ["ParsedOutput", "NVARCHAR(MAX)", "NULL", "Kết quả đã phân tích thành cấu trúc."],
        ["AgentVersion", "VARCHAR(16)", "NULL", "Phiên bản khuôn mẫu lời nhắc (prompt)."],
        ["ModelId", "NVARCHAR(100)", "NULL", "Tên mô hình đã gọi."],
        ["ParseSuccess", "BIT", "NULL", "Phân tích kết quả thành công hay không."],
        ["LatencyMs", "INT", "NULL", "Thời gian gọi mô hình (mili-giây)."],
    ])

    d.h3("3.2.4.", "Nhóm bảng Kết quả phân tích")
    d.para("Nhóm bảng này lưu các kết quả trung gian ở mức thành phần và kết quả cuối ở mức "
           "công trình. Trong đó, hai bảng quan trọng nhất của quy trình hiện hành là "
           "BuildingStyleHypotheses (các giả thuyết phong cách từ hội đồng giám khảo) và "
           "BuildingStyleResults (kết quả phong cách cuối cùng kèm bản ghi đầy đủ).")
    _dd(d, "Cấu trúc bảng GeometricFeatures (đặc trưng hình học của thành phần)", [
        ["FeatureId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh đặc trưng."],
        ["ComponentId", "UNIQUEIDENTIFIER", "Khóa ngoại → Components", "Thành phần tương ứng."],
        ["Description", "NVARCHAR(MAX)", "NULL", "Mô tả đặc trưng quan sát được."],
        ["StructuredData", "NVARCHAR(MAX)", "NULL", "Đặc trưng ở dạng có cấu trúc (JSON)."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
    ])
    _dd(d, "Cấu trúc bảng StylePredictions (phân bố phong cách của thành phần)", [
        ["PredictionId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh dự đoán."],
        ["ComponentId", "UNIQUEIDENTIFIER", "Khóa ngoại → Components", "Thành phần tương ứng."],
        ["StyleName", "NVARCHAR(100)", "NULL", "Tên phong cách dự đoán."],
        ["Probability", "FLOAT", "NULL", "Trọng số/điểm xếp hạng của phong cách."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
    ])
    _dd(d, "Cấu trúc bảng ConsistencyChecks (kiểm tra nhất quán)", [
        ["CheckId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh lần kiểm tra."],
        ["ComponentId", "UNIQUEIDENTIFIER", "Khóa ngoại → Components", "Thành phần tương ứng."],
        ["IsConsistent", "BIT", "NULL", "Kết quả nhất quán hay mâu thuẫn."],
        ["Reason", "NVARCHAR(MAX)", "NULL", "Diễn giải lý do."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
    ])
    _dd(d, "Cấu trúc bảng ComponentFinalStyles (phong cách cuối của thành phần)", [
        ["Id", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh bản ghi."],
        ["ComponentId", "UNIQUEIDENTIFIER", "Khóa ngoại → Components", "Thành phần tương ứng."],
        ["FinalStyle", "NVARCHAR(100)", "NULL", "Phong cách kết luận cho thành phần."],
        ["Confidence", "FLOAT", "NULL", "Độ tin cậy."],
        ["Explanation", "NVARCHAR(MAX)", "NULL", "Giải thích."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
    ])
    _dd(d, "Cấu trúc bảng BuildingStyleHypotheses (giả thuyết phong cách công trình)", [
        ["HypothesisId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh giả thuyết."],
        ["ImageId", "UNIQUEIDENTIFIER", "Khóa ngoại → Images", "Ảnh công trình."],
        ["AgentId", "UNIQUEIDENTIFIER", "Khóa ngoại → Agents", "Giám khảo đưa ra giả thuyết."],
        ["HypothesisType", "NVARCHAR(50)", "NULL", "Loại giả thuyết (ví dụ giám khảo, phương án)."],
        ["StyleName", "NVARCHAR(100)", "NULL", "Phong cách đề xuất."],
        ["Confidence", "FLOAT", "NULL", "Độ tin cậy do giám khảo gán."],
        ["Explanation", "NVARCHAR(MAX)", "NULL", "Lập luận của giám khảo."],
        ["BasedOnFeatures", "NVARCHAR(MAX)", "NULL", "Bằng chứng làm cơ sở (JSON)."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
    ])
    _dd(d, "Cấu trúc bảng BuildingStyleResults (kết quả phong cách công trình)", [
        ["ResultId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh kết quả."],
        ["ImageId", "UNIQUEIDENTIFIER", "Khóa ngoại → Images", "Ảnh công trình."],
        ["KeyEvidence", "NVARCHAR(MAX)", "NULL", "Các bằng chứng then chốt (JSON)."],
        ["FinalStyle", "NVARCHAR(100)", "NULL", "Phong cách kết luận cuối cùng."],
        ["Confidence", "FLOAT", "NULL", "Độ tin cậy báo cáo."],
        ["Explanation", "NVARCHAR(MAX)", "NULL", "Giải thích kết luận."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm tạo."],
        ["DetailJson", "NVARCHAR(MAX)", "NULL", "Toàn bộ kết quả phân tích (JSON) để mở lại."],
    ])

    d.h3("3.2.5.", "Nhóm bảng Nhật ký hệ thống")
    _dd(d, "Cấu trúc bảng SystemLogs (nhật ký hệ thống)", [
        ["LogId", "UNIQUEIDENTIFIER", "Khóa chính", "Định danh dòng nhật ký."],
        ["LogLevel", "NVARCHAR(50)", "NULL", "Mức độ: INFO | WARNING | ERROR."],
        ["Message", "NVARCHAR(MAX)", "NULL", "Nội dung nhật ký."],
        ["CreatedAt", "DATETIME", "Mặc định GETDATE()", "Thời điểm ghi nhận."],
    ])

    d.h2("3.3.", "Quan hệ khóa ngoại và ràng buộc toàn vẹn")
    d.para("Các quan hệ khóa ngoại bảo đảm toàn vẹn tham chiếu giữa các bảng. Để việc xóa "
           "dữ liệu nhất quán, các quan hệ chính được thiết lập cơ chế xóa lan truyền "
           "(ON DELETE CASCADE): khi xóa một dự án thì toàn bộ ảnh thuộc dự án bị xóa theo; "
           "khi xóa một ảnh thì các thành phần, lần chạy tác tử, giả thuyết và kết quả "
           "phong cách liên quan cũng bị xóa. Riêng quan hệ giữa AgentRuns và Components "
           "không đặt khóa ngoại lan truyền nhằm tránh xung đột nhiều đường xóa lan truyền "
           "(multiple cascade paths) mà SQL Server không cho phép. Ngoài ra, bảng AgentRuns "
           "có ràng buộc kiểm tra CK_AgentRuns_Target bảo đảm mỗi lần chạy phải gắn với ít "
           "nhất một trong hai khóa ImageId hoặc ComponentId.")

    d.h2("3.4.", "Đánh giá chuẩn hóa dữ liệu")
    d.para("Lược đồ cơ sở dữ liệu được thiết kế đạt đến dạng chuẩn 3 (3NF). Cụ thể, mỗi "
           "bảng đều có khóa chính rõ ràng và các thuộc tính nguyên tử (dạng chuẩn 1 — "
           "1NF); mọi thuộc tính không khóa phụ thuộc đầy đủ vào toàn bộ khóa chính (dạng "
           "chuẩn 2 — 2NF) do khóa chính là một định danh đơn; và không tồn tại phụ thuộc "
           "bắc cầu giữa các thuộc tính không khóa (dạng chuẩn 3 — 3NF).")
    d.para("Một số đánh đổi thiết kế được chấp nhận có chủ đích. Các trường dạng chuỗi JSON "
           "(như BoundingBox, KeyEvidence, DetailJson, BasedOnFeatures) lưu dữ liệu bán cấu "
           "trúc nhằm giữ lại đầy đủ kết quả phân tích để mở lại về sau, đổi lấy việc không "
           "truy vấn sâu vào nội dung bên trong. Việc lưu DetailJson giúp một phân tích cũ "
           "có thể được tái hiện đầy đủ mà không cần chạy lại quy trình tốn kém.")

    d.h2("3.5.", "Kết chương")
    d.para("Chương 3 đã trình bày thiết kế cơ sở dữ liệu của hệ thống gồm 14 bảng thuộc năm "
           "nhóm chức năng, mô hình thực thể – quan hệ tổng thể, mô tả chi tiết từng bảng, "
           "các ràng buộc khóa ngoại và mức độ chuẩn hóa. Đây là nền tảng lưu trữ cho các "
           "chức năng được phân tích và thiết kế trong Chương 4.")


# ═══════════════════════ CHƯƠNG 4 — PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG ═════════════
def chuong_4(d):
    d.chapter_heading("Phân tích và thiết kế hệ thống")

    d.h2("4.1.", "Khảo sát hiện trạng và tổng quan hệ thống")
    d.h3("4.1.1.", "Mô tả tổng quan hệ thống")
    d.para("Hệ thống là một ứng dụng web cho phép người dùng tải lên ảnh chụp một công "
           "trình kiến trúc và nhận về phong cách kiến trúc của công trình kèm lời giải "
           "thích. Khác với các công cụ phân loại đóng chỉ nhận biết một số ít phong cách "
           "cố định, hệ thống hướng tới nhận dạng mở từ vựng dựa trên cơ sở tri thức 106 "
           "phong cách và sự phối hợp của nhiều mô hình thị giác – ngôn ngữ. Hệ thống còn "
           "lưu lịch sử phân tích, hỗ trợ hỏi đáp tiếp theo về kết quả, hiển thị song ngữ "
           "Việt – Anh và cung cấp trang quản trị cho quản trị viên.")
    d.h3("4.1.2.", "Đối tượng sử dụng")
    d.para("Hệ thống phục vụ hai nhóm tác nhân chính. Người dùng đăng nhập bằng tài khoản "
           "Google để tải ảnh, xem kết quả phân tích, tra cứu lịch sử và đặt câu hỏi tiếp "
           "theo. Quản trị viên quản lý người dùng (kích hoạt/khóa tài khoản), quản lý dự "
           "án và ảnh, theo dõi nhật ký, danh sách tác tử và các thống kê hoạt động.")

    d.h2("4.2.", "Yêu cầu hệ thống")
    d.h3("4.2.1.", "Yêu cầu chức năng")
    d.bullet("Đăng nhập, đăng xuất bằng tài khoản Google; lấy thông tin hồ sơ người dùng.")
    d.bullet("Tải ảnh công trình lên hệ thống với kiểm tra định dạng và dung lượng.")
    d.bullet("Phân tích ảnh để xác định phong cách kiến trúc kèm phân bố phong cách, bằng "
             "chứng và giải thích song ngữ.")
    d.bullet("Từ chối trả lời khi mức đồng thuận thấp; báo độ tin cậy của kết quả.")
    d.bullet("Lưu và tra cứu lịch sử phân tích; mở lại kết quả cũ; hỏi đáp tiếp theo dựa "
             "trên bằng chứng của một phân tích.")
    d.bullet("Tra cứu cơ sở tri thức phong cách (thẻ thông tin phong cách).")
    d.bullet("Quản trị: quản lý người dùng, dự án, ảnh; xem nhật ký, tác tử và thống kê.")
    d.h3("4.2.2.", "Yêu cầu phi chức năng")
    d.bullet("Độ tin cậy: chịu lỗi khi một nhà cung cấp mô hình tạm gián đoạn nhờ cơ chế "
             "thử lại và dự phòng; không sập toàn hệ thống khi một giám khảo gặp lỗi.")
    d.bullet("Bảo mật: xác thực bằng JWT, phân quyền theo vai trò, khóa tài khoản bị vô "
             "hiệu hóa, đặt tên tệp theo định danh duy nhất để tránh tấn công đường dẫn.")
    d.bullet("Hiệu năng: xử lý một ảnh trong khoảng thời gian chấp nhận được (trung bình "
             "khoảng 10–11 giây/ảnh trong thử nghiệm).")
    d.bullet("Khả năng mở rộng: thêm phong cách mới bằng cách bổ sung tri thức, không cần "
             "huấn luyện lại; dễ thay đổi hoặc bổ sung nhà cung cấp mô hình.")
    d.bullet("Khả dụng và trải nghiệm: giao diện song ngữ, chế độ sáng/tối, phản hồi rõ "
             "ràng khi hệ thống không chắc chắn.")

    d.h2("4.3.", "Kiến trúc tổng thể hệ thống")
    d.para("Hệ thống được tổ chức theo kiến trúc nhiều tầng. Tầng giao diện (frontend) viết "
           "bằng React chạy trên trình duyệt. Tầng thân (backend) viết bằng FastAPI cung "
           "cấp các giao diện lập trình ứng dụng theo kiến trúc REST, điều phối quy trình "
           "phân tích và truy xuất cơ sở dữ liệu. Tầng dữ liệu là SQL Server. Quy trình "
           "phân tích gọi đến các dịch vụ mô hình ngôn ngữ – thị giác bên ngoài (Gemini, "
           "OpenAI, Grok, DeepSeek) qua API. Kiến trúc tổng thể được minh họa trong Hình 4.1.")
    d.figure("Sơ đồ kiến trúc tổng thể của hệ thống")

    d.h2("4.4.", "Quy trình phân tích đa tác tử mở từ vựng")
    d.h3("4.4.1.", "Tổng quan luồng xử lý")
    d.para("Trọng tâm của hệ thống là quy trình phân tích đa tác tử mở từ vựng. Khi nhận "
           "một ảnh, quy trình lần lượt đi qua các bước: trích xuất bằng chứng từ ảnh "
           "(Tác tử A) → nối khớp cơ sở tri thức để sinh tập phong cách ứng viên → hội "
           "đồng ba giám khảo thị giác chấm điểm độc lập → đo mức đồng thuận và gộp phân "
           "bố → trọng tài đưa ra kết luận → cơ chế từ chối khi không đủ chắc chắn → dịch "
           "song ngữ kết quả. Toàn bộ luồng được minh họa trong Hình 4.2.")
    d.figure("Sơ đồ luồng quy trình phân tích đa tác tử")

    d.h3("4.4.2.", "Tác tử A — Trích xuất bằng chứng")
    d.para("Tác tử A sử dụng mô hình thị giác – ngôn ngữ để “đọc” ảnh và điền một phiếu "
           "bằng chứng có cấu trúc gồm khoảng 12 chiều quan sát: khối hình tổng thể, mái, "
           "hệ đỡ/cột, dạng vòm, cửa, mặt đứng, trang trí, vật liệu, tính thẳng đứng, vòm "
           "trần/mái vòm, tổ chức không gian và chi tiết đặc trưng. Để giảm nhiễu ngẫu "
           "nhiên, bước trích xuất được gọi nhiều lượt trên nhiều nhà cung cấp; mỗi phong "
           "cách do mô hình đề xuất được bỏ phiếu theo định danh trong cơ sở tri thức, qua "
           "đó gộp các cách gọi đồng nghĩa và giữ lại các đề xuất ổn định.")

    d.h3("4.4.3.", "Nối khớp cơ sở tri thức và sinh tập ứng viên")
    d.para("Cơ sở tri thức gồm 106 phong cách thuộc 12 họ, mỗi mục có tên chuẩn, tên đồng "
           "nghĩa, họ, đặc trưng nhận dạng và nguồn tham chiếu. Tên phong cách do Tác tử A "
           "đề xuất được nối khớp về cơ sở tri thức theo ba mức (khớp chính xác, khớp đồng "
           "nghĩa, khớp xấp xỉ với ngưỡng tương đồng 0,82), sau đó hệ thống sinh ra một tập "
           "phong cách ứng viên (top-K) gồm các phong cách được bỏ phiếu cùng các phong "
           "cách lân cận cùng họ nhằm bảo đảm độ thu hồi (recall). Tập ứng viên này giới "
           "hạn không gian lựa chọn cho hội đồng giám khảo.")

    d.h3("4.4.4.", "Hội đồng ba giám khảo thị giác")
    d.para("Hội đồng gồm ba giám khảo là ba mô hình thị giác – ngôn ngữ độc lập của ba nhà "
           "cung cấp khác nhau (Gemini, OpenAI và Grok). Mỗi giám khảo cùng nhìn ảnh và "
           "phiếu bằng chứng, rồi chấm một phân bố xác suất trên tập phong cách ứng viên. "
           "Việc dùng ba nhà cung cấp độc lập giúp các lỗi ít tương quan với nhau, làm cho "
           "việc cân bằng trọng số giữa các giám khảo trở nên chính đáng. Nếu một giám "
           "khảo gặp lỗi, hội đồng vẫn hoạt động với các giám khảo còn lại.")

    d.h3("4.4.5.", "Đo đồng thuận và gộp phân bố")
    d.para("Mức đồng thuận giữa các giám khảo được đo bằng tương quan thứ hạng Spearman "
           "trung bình theo từng cặp giám khảo, thay vì so sánh trực tiếp các giá trị xác "
           "suất (vì mỗi mô hình có mức “rộng tay” khác nhau nhưng thứ hạng thì so sánh "
           "được). Các phân bố sau đó được gộp lại có trọng số theo mức quyết đoán của từng "
           "giám khảo để tạo thành một phân bố tổng hợp.")

    d.h3("4.4.6.", "Trọng tài và các cơ chế bổ trợ")
    d.para("Trọng tài là mô hình mạnh (GPT-4o) nhìn ảnh đầy đủ cùng các phiếu của hội đồng "
           "và điểm đồng thuận để đưa ra kết luận theo lối chuỗi suy luận. Trọng tài được "
           "trang bị ba cơ chế bổ trợ: (1) cơ chế gỡ trần thu hồi (escape-hatch) cho phép "
           "nêu một phong cách ngoài tập ứng viên trong một lượt, với điều kiện phong cách "
           "đó vừa khớp cơ sở tri thức vừa được đặc trưng quan sát ủng hộ; (2) cơ chế phân "
           "định tương phản (run-off) khi hai phong cách dẫn đầu sát nhau, dùng đặc trưng "
           "nhận dạng của cơ sở tri thức làm danh mục đối chiếu; (3) cơ chế “đọc tự do” "
           "(free-read) cho một mô hình gọi tên phong cách trực tiếp từ ảnh như một kênh "
           "độc lập bổ sung. Khi trọng tài gặp lỗi, hệ thống chuyển sang nhà cung cấp dự "
           "phòng để bảo đảm vẫn có kết luận.")

    d.h3("4.4.7.", "Cơ chế từ chối trả lời")
    d.para("Hệ thống áp dụng cơ chế từ chối (abstention): khi mức đồng thuận giữa các giám "
           "khảo thấp hơn ngưỡng, hoặc số giám khảo hợp lệ ít, hoặc khoảng cách giữa hai "
           "phong cách dẫn đầu quá nhỏ và độ phân tán cao, hệ thống đánh dấu kết quả là "
           "“không chắc chắn”. Việc này giúp hệ thống tránh kết luận sai một cách tự tin, "
           "đặc biệt với các công trình mang tính lai ghép nhiều phong cách.")

    d.h3("4.4.8.", "Dịch song ngữ và báo cáo độ tin cậy")
    d.para("Phần giải thích và bằng chứng được dịch sang tiếng Việt bằng mô hình DeepSeek, "
           "tách thành nhiều lời gọi song song để tránh bị cắt cụt nội dung. Độ tin cậy "
           "báo cáo cho người dùng được tính bằng tích của độ sắc nét của phân bố và một "
           "hệ số chất lượng lần chạy (tổng hợp từ độ đầy đủ của giám khảo, độ đầy đủ của "
           "trích xuất và mức đồng thuận), nhằm phản ánh trung thực mức độ chắc chắn thay "
           "vì chỉ dựa vào con số mô hình tự gán.")

    d.h2("4.5.", "Phân tích ca sử dụng")
    d.para("Hệ thống có hai tác nhân chính là người dùng và quản trị viên. Sơ đồ ca sử "
           "dụng tổng thể (Hình 4.3) thể hiện các nhóm chức năng theo từng tác nhân.")
    d.figure("Sơ đồ ca sử dụng tổng thể của hệ thống")
    d.figure("Sơ đồ ca sử dụng của người dùng")
    d.figure("Sơ đồ ca sử dụng của quản trị viên")

    d.h2("4.6.", "Đặc tả sơ đồ hoạt động")
    d.para("Các sơ đồ hoạt động mô tả luồng xử lý của những quy trình nghiệp vụ chính: "
           "đăng nhập, tải ảnh và phân tích, cơ chế từ chối, và quản trị người dùng.")
    d.figure("Sơ đồ hoạt động chức năng đăng nhập bằng Google")
    d.figure("Sơ đồ hoạt động chức năng tải ảnh và phân tích phong cách")
    d.figure("Sơ đồ hoạt động cơ chế từ chối khi không đủ chắc chắn")

    d.h2("4.7.", "Đặc tả sơ đồ trình tự")
    d.para("Các sơ đồ trình tự mô tả thứ tự tương tác giữa các thành phần theo thời gian "
           "cho các luồng chính của hệ thống.")
    d.figure("Sơ đồ trình tự chức năng đăng nhập bằng Google OAuth")
    d.figure("Sơ đồ trình tự quy trình phân tích ảnh qua các tác tử")
    d.figure("Sơ đồ trình tự chức năng quản trị người dùng")

    d.h2("4.8.", "Thiết kế sơ đồ lớp, gói và thành phần")
    d.para("Sơ đồ lớp mô tả cấu trúc tĩnh của các lớp dịch vụ và mô hình dữ liệu; sơ đồ "
           "gói thể hiện cách tổ chức mã nguồn theo các gói chức năng (app, chatbot, "
           "frontend); sơ đồ thành phần thể hiện các thành phần phần mềm và quan hệ phụ "
           "thuộc giữa chúng.")
    d.figure("Sơ đồ lớp các dịch vụ của quy trình phân tích")
    d.figure("Sơ đồ gói tổ chức mã nguồn hệ thống")
    d.figure("Sơ đồ thành phần của hệ thống")

    d.h2("4.9.", "Sơ đồ triển khai")
    d.para("Sơ đồ triển khai (Hình 4.15) mô tả cách phân bố các thành phần trên hạ tầng "
           "vận hành: trình duyệt người dùng, máy chủ ứng dụng (frontend và backend), máy "
           "chủ cơ sở dữ liệu SQL Server và các dịch vụ mô hình ngôn ngữ – thị giác bên "
           "ngoài.")
    d.figure("Sơ đồ triển khai hệ thống")

    d.h2("4.10.", "Thiết kế giao diện lập trình ứng dụng")
    d.para("Phần thân hệ thống cung cấp các điểm cuối (endpoint) theo kiến trúc REST. Bảng "
           "4.1 liệt kê các điểm cuối chính cùng yêu cầu xác thực và mô tả.")
    d.table_caption("Danh sách các điểm cuối API chính của hệ thống")
    d.data_table(
        ["Phương thức", "Đường dẫn", "Xác thực", "Mô tả"],
        [
            ["GET", "/health", "Không", "Kiểm tra tình trạng hệ thống."],
            ["GET", "/auth/google/login", "Không", "Chuyển hướng đăng nhập Google."],
            ["POST", "/auth/google", "Không", "Đổi mã Google lấy JWT của hệ thống."],
            ["GET", "/auth/me", "JWT", "Lấy hồ sơ người dùng hiện tại."],
            ["POST", "/upload/image", "JWT", "Tải ảnh lên, trả về định danh tệp."],
            ["POST", "/analyze", "JWT", "Chạy quy trình phân tích, trả kết quả."],
            ["GET", "/analyze/history", "JWT", "Lấy lịch sử phân tích của người dùng."],
            ["GET", "/analyze/history/{id}", "JWT", "Mở lại một phân tích đã lưu."],
            ["GET", "/analyze/image/{id}", "JWT", "Tải lại ảnh gốc của một phân tích."],
            ["POST", "/analyze/{id}/ask", "JWT", "Hỏi đáp tiếp theo dựa trên bằng chứng."],
            ["GET", "/knowledge/style", "JWT", "Tra cứu thẻ thông tin phong cách."],
            ["GET, PATCH", "/admin/users", "JWT (admin)", "Quản lý người dùng."],
            ["GET, DELETE", "/admin/projects, /admin/images", "JWT (admin)", "Quản lý dự án, ảnh."],
            ["GET", "/admin/logs, /admin/agents, /admin/stats", "JWT (admin)", "Nhật ký, tác tử, thống kê."],
        ],
        widths=[2.2, 5.2, 2.4, 6.2],
    )

    d.h2("4.11.", "Kết chương")
    d.para("Chương 4 đã trình bày yêu cầu hệ thống, kiến trúc tổng thể và đặc biệt là quy "
           "trình phân tích đa tác tử mở từ vựng — đóng góp cốt lõi của đề tài — cùng các "
           "sơ đồ UML và thiết kế giao diện lập trình ứng dụng. Chương 5 sẽ trình bày môi "
           "trường triển khai và kết quả thử nghiệm thực tế của hệ thống.")


# ═══════════════════ CHƯƠNG 5 — MÔI TRƯỜNG VÀ KẾT QUẢ THỬ NGHIỆM ═══════════════
def chuong_5(d):
    d.chapter_heading("Môi trường và kết quả thử nghiệm")

    d.h2("5.1.", "Môi trường triển khai và thử nghiệm")
    d.h3("5.1.1.", "Môi trường phần cứng và phần mềm")
    d.para("Hệ thống được phát triển và thử nghiệm trên một máy tính cá nhân với cấu hình "
           "trong Bảng 5.1. Vì phần xử lý trí tuệ nhân tạo được thực hiện bởi các dịch vụ "
           "mô hình bên ngoài qua API, hệ thống không yêu cầu cấu hình phần cứng cao.")
    d.table_caption("Cấu hình phần cứng môi trường thử nghiệm")
    d.data_table(
        ["Thành phần", "Thông số"],
        [
            ["Bộ xử lý (CPU)", "Intel Core i5-12500H (2,50 GHz)"],
            ["Bộ nhớ (RAM)", "16 GB"],
            ["Card đồ họa (GPU)", "NVIDIA GeForce RTX 3050 Laptop (4 GB)"],
            ["Ổ lưu trữ", "SSD (477 GB)"],
            ["Hệ điều hành", "Windows 11, 64-bit (x64)"],
        ],
        widths=[6.0, 9.0],
    )
    d.para("Về phần mềm, phần thân hệ thống được viết bằng Python 3.10 với FastAPI và "
           "SQLAlchemy bất đồng bộ; cơ sở dữ liệu là SQL Server; phần giao diện viết bằng "
           "React và Vite. Quy trình phân tích sử dụng dịch vụ của bốn nhà cung cấp mô "
           "hình: Gemini, OpenAI, Grok và DeepSeek, được gọi qua API.")
    d.h3("5.1.2.", "Phạm vi dữ liệu thử nghiệm")
    d.para("Bộ dữ liệu thử nghiệm gồm 200 ảnh tương ứng 100 phong cách kiến trúc, mỗi phong "
           "cách có 2 ảnh được sinh từ hai nguồn khác nhau (một ảnh sinh bởi ChatGPT và "
           "một ảnh sinh bởi Gemini). Mỗi ảnh có nhãn phong cách chính làm đáp án tham "
           "chiếu. Việc dùng hai nguồn sinh ảnh khác nhau giúp đánh giá tính khách quan, "
           "tránh thiên lệch về một nguồn ảnh cụ thể.")

    d.h2("5.2.", "Phương pháp đánh giá")
    d.h3("5.2.1.", "Quy trình chấm điểm tự động")
    d.para("Hệ thống được đánh giá bằng một quy trình chấm điểm tự động. Một dự đoán được "
           "tính là đúng (top-1) khi định danh phong cách dự đoán trùng với định danh phong "
           "cách đáp án sau khi chuẩn hóa và nối khớp về cơ sở tri thức (gộp các tên đồng "
           "nghĩa). Để bảo đảm tính công bằng, hai mô hình nền (Gemini và ChatGPT) cũng "
           "được chấm thủ công trên cùng thang đo và cùng bộ ảnh. Những ảnh mà mô hình gặp "
           "lỗi gọi (không cho được tên) được đánh dấu lỗi và loại khỏi mẫu, không tính là "
           "sai.")
    d.h3("5.2.2.", "Các chỉ số đo lường")
    d.bullet("Độ thu hồi ứng viên (recall@K): tỉ lệ ảnh mà phong cách đáp án có mặt trong "
             "tập ứng viên — đây là trần cứng cho độ chính xác.")
    d.bullet("Độ chính xác top-1 và top-3: tỉ lệ ảnh mà phong cách đáp án là dự đoán đứng "
             "đầu, hoặc nằm trong ba phong cách đứng đầu.")
    d.bullet("Độ chính xác theo họ (family accuracy): tỉ lệ ảnh mà dự đoán cùng họ phong "
             "cách với đáp án (đúng họ nhưng có thể sai phong cách con).")
    d.bullet("Tỉ lệ mắc bẫy nhiễu (distractor rate): tỉ lệ ảnh hệ thống chọn nhầm một đặc "
             "trưng nhiễu cố ý — càng thấp càng tốt.")
    d.bullet("Tỉ lệ chuyển hóa (conversion): tỉ lệ chọn đúng top-1 khi phong cách đáp án đã "
             "có trong tập ứng viên — phản ánh chất lượng khâu chấm điểm.")
    d.bullet("Đường đánh đổi rủi ro – độ phủ (risk-coverage): quan hệ giữa tỉ lệ ảnh được "
             "trả lời (độ phủ) và độ chính xác khi áp ngưỡng từ chối khác nhau.")

    d.h2("5.3.", "Kết quả định lượng")
    d.h3("5.3.1.", "Kết quả tổng thể")
    d.para("Trên toàn bộ 200 ảnh (không có ảnh lỗi), hệ thống đạt các chỉ số trong Bảng 5.2. "
           "Độ thu hồi ứng viên đạt 93,0% cho thấy phần lớn đáp án đã được đưa vào tập ứng "
           "viên; độ chính xác top-1 đạt 74,5% và top-3 đạt 88,5%; độ chính xác theo họ đạt "
           "84,0%. Tỉ lệ mắc bẫy nhiễu bằng 0%, cho thấy hệ thống không bị đánh lừa bởi các "
           "đặc trưng nhiễu cố ý.")
    d.table_caption("Kết quả các chỉ số đánh giá tổng thể của hệ thống (200 ảnh)")
    d.data_table(
        ["Chỉ số", "Giá trị"],
        [
            ["Độ thu hồi ứng viên (recall@K)", "93,0% (186/200)"],
            ["Độ chính xác top-1", "74,5% (149/200)"],
            ["Độ chính xác top-3", "88,5% (177/200)"],
            ["Độ chính xác theo họ", "84,0% (168/200)"],
            ["Tỉ lệ chuyển hóa (conversion)", "80,1% (149/186)"],
            ["Tỉ lệ mắc bẫy nhiễu (distractor)", "0,0%"],
            ["Tỉ lệ từ chối (abstain)", "3,5% (7/200)"],
            ["Thời gian xử lý trung bình", "10,6 giây/ảnh"],
        ],
        widths=[8.0, 7.0],
    )

    d.h3("5.3.2.", "Đối sánh với hai mô hình nền")
    d.para("Để đánh giá khách quan, hệ thống được đối sánh với hai mô hình nền là Gemini và "
           "ChatGPT trong vai trò bộ phân loại trực tiếp, trên cùng 200 ảnh và cùng thang "
           "đo top-1. Bảng 5.3 trình bày kết quả tách theo nguồn sinh ảnh.")
    d.table_caption("Đối sánh độ chính xác top-1 giữa hệ thống và hai mô hình nền")
    d.data_table(
        ["Phương pháp", "Ảnh nguồn ChatGPT", "Ảnh nguồn Gemini", "Tổng (200 ảnh)"],
        [
            ["Gemini (bộ phân loại)", "78%", "66%", "72,0%"],
            ["ChatGPT (bộ phân loại)", "84%", "65%", "74,5%"],
            ["Hệ thống (đề xuất)", "78%", "71%", "74,5%"],
        ],
        widths=[5.0, 3.6, 3.6, 3.6],
    )
    d.para("Trên cùng một thang đo nghiêm ngặt, hệ thống đạt 74,5%, ngang bằng mô hình nền "
           "mạnh hơn là ChatGPT (74,5%) và vượt Gemini (72,0%). Đáng chú ý, mỗi mô hình nền "
           "đạt độ chính xác cao nhất trên chính những ảnh do nó sinh ra (Gemini 66% trên "
           "ảnh-Gemini nhưng ChatGPT đạt 84% trên ảnh-ChatGPT), cho thấy thiên kiến “sân "
           "nhà”. Trong khi đó, hệ thống cân bằng hơn giữa hai nguồn ảnh (78% và 71%) nhờ "
           "phối hợp nhiều nhà cung cấp độc lập, nên ít lệ thuộc vào nguồn sinh ảnh.")

    d.h3("5.3.3.", "Phân tích theo độ phủ và mức đồng thuận")
    d.para("Cơ chế từ chối cho phép đánh đổi giữa độ phủ và độ chính xác: khi nâng ngưỡng "
           "đồng thuận, hệ thống trả lời ít ảnh hơn nhưng giữ độ chính xác trên phần đã trả "
           "lời. Trong thử nghiệm, ở ngưỡng đồng thuận 0,5 hệ thống vẫn trả lời 97,5% số "
           "ảnh; tỉ lệ từ chối tổng thể chỉ 3,5%, cho thấy hệ thống thận trọng đúng mức mà "
           "không từ chối quá nhiều. Đường đánh đổi rủi ro – độ phủ được minh họa trong "
           "Hình 5.1.")
    d.figure("Đường đánh đổi rủi ro – độ phủ theo ngưỡng đồng thuận")

    d.h3("5.3.4.", "So sánh với các công trình liên quan")
    d.para("So với các công trình trước, hướng phân loại đóng của Xu và cộng sự (2014) đạt "
           "khoảng 46% trên 25 phong cách và độ chính xác giảm khi tăng số lớp. Công trình "
           "ArchiLense (2025) dựa trên mô hình thị giác – ngôn ngữ đạt độ chính xác phân "
           "loại khoảng 84,5%. Hệ thống của đề tài đạt độ chính xác theo họ 84,0% trên một "
           "tập phong cách rộng hơn nhiều (106 phong cách, mở từ vựng) mà không cần huấn "
           "luyện, đồng thời bổ sung cơ chế từ chối và giải thích bám bằng chứng — những "
           "yếu tố mà các công trình trước thường không có.")

    d.h2("5.4.", "Kết quả giao diện hệ thống")
    d.para("Phần này trình bày giao diện thực tế của hệ thống. Các hình minh họa sẽ được "
           "bổ sung trong phiên bản hoàn chỉnh.")
    d.figure("Giao diện trang chủ (Landing) của hệ thống")
    d.figure("Giao diện đăng nhập bằng tài khoản Google")
    d.figure("Giao diện tải ảnh và kết quả phân tích phong cách của người dùng")
    d.figure("Giao diện khối bằng chứng và phân bố phong cách")
    d.figure("Giao diện lịch sử phân tích của người dùng")
    d.figure("Giao diện bảng điều khiển quản trị viên")

    d.h2("5.5.", "Kết chương")
    d.para("Chương 5 đã trình bày môi trường triển khai, phương pháp đánh giá và kết quả "
           "thử nghiệm định lượng của hệ thống. Trên 200 ảnh, hệ thống đạt độ chính xác "
           "top-1 74,5% — ngang bằng mô hình nền mạnh và vượt mô hình nền còn lại trên cùng "
           "thang đo — đồng thời cân bằng hơn giữa các nguồn ảnh và bổ sung các giá trị về "
           "khả năng mở rộng, từ chối và giải thích. Chương 6 tổng kết kết quả đạt được và "
           "hướng phát triển.")


# ═══════════════ CHƯƠNG 6 — KẾT QUẢ VÀ HƯỚNG PHÁT TRIỂN CỦA ĐỀ TÀI ═════════════
def chuong_6(d):
    d.chapter_heading("Kết quả và hướng phát triển của đề tài")

    d.h2("6.1.", "Kết quả đạt được")
    d.para("Đề tài đã hoàn thành các mục tiêu đề ra và đạt được những kết quả chính sau:")
    d.bullet("Xây dựng cơ sở tri thức phong cách kiến trúc gồm 106 phong cách thuộc 12 họ, "
             "có đặc trưng nhận dạng và nguồn tham chiếu, làm nền tảng cho nhận dạng mở từ "
             "vựng.")
    d.bullet("Thiết kế và hiện thực quy trình phân tích đa tác tử mở từ vựng: trích xuất "
             "bằng chứng, nối khớp tri thức, hội đồng ba giám khảo thị giác độc lập, đo "
             "đồng thuận và trọng tài, kèm các cơ chế bổ trợ và cơ chế từ chối.")
    d.bullet("Xây dựng ứng dụng web hoàn chỉnh, đa vai trò: xác thực bằng Google, trang "
             "phân tích và lịch sử cho người dùng, hỏi đáp tiếp theo, trang quản trị, giao "
             "diện song ngữ Việt – Anh.")
    d.bullet("Thiết lập cơ sở dữ liệu 14 bảng trên SQL Server lưu trữ đầy đủ người dùng, dự "
             "án, ảnh, các lần chạy tác tử và kết quả phân tích.")
    d.bullet("Đánh giá định lượng trên 200 ảnh: độ chính xác top-1 đạt 74,5%, ngang bằng "
             "ChatGPT và vượt Gemini trên cùng thang đo, với độ cân bằng tốt hơn giữa các "
             "nguồn ảnh; độ chính xác theo họ đạt 84,0%.")

    d.h2("6.2.", "Hạn chế của đề tài")
    d.para("Bên cạnh các kết quả đạt được, đề tài còn một số hạn chế:")
    d.bullet("Hệ thống phụ thuộc vào các dịch vụ mô hình bên ngoài, nên chi phí gọi API và "
             "thời gian xử lý (trung bình khoảng 10–11 giây/ảnh) còn đáng kể và chịu ảnh "
             "hưởng bởi tình trạng dịch vụ.")
    d.bullet("Bộ dữ liệu thử nghiệm là ảnh sinh bởi mô hình; cần mở rộng đánh giá trên ảnh "
             "chụp thực tế đa dạng hơn để khẳng định khả năng tổng quát hóa.")
    d.bullet("Với một số cặp phong cách rất gần nhau (ví dụ phong cách gốc và biến thể phục "
             "hưng của nó), hệ thống vẫn còn nhầm lẫn, làm hạn chế độ chính xác top-1 theo "
             "thang đo nghiêm ngặt.")
    d.bullet("Cơ sở tri thức còn cần bổ sung định danh tham chiếu chuẩn (mã AAT, Wikidata) "
             "đầy đủ cho từng phong cách.")

    d.h2("6.3.", "Hướng phát triển")
    d.para("Trên cơ sở các hạn chế, đề tài đề xuất các hướng phát triển tiếp theo:")
    d.bullet("Bổ sung một kênh “đọc tự do” thứ hai và để trọng tài lựa chọn giữa các cách "
             "đọc độc lập, nhằm nâng độ chính xác trên các cặp phong cách dễ nhầm.")
    d.bullet("Mở rộng và làm giàu cơ sở tri thức (thêm phong cách, đặc trưng, hình mẫu và "
             "định danh tham chiếu chuẩn).")
    d.bullet("Đánh giá trên tập ảnh chụp thực tế quy mô lớn, đa vùng miền; tối ưu chi phí "
             "và thời gian xử lý (đệm kết quả, gọi song song hợp lý).")
    d.bullet("Phát triển ứng dụng di động và bổ sung bản đồ phân bố phong cách theo địa "
             "điểm, hướng tới ứng dụng du lịch và giáo dục di sản.")

    d.h2("6.4.", "Kết luận chung")
    d.para("Đề tài đã xây dựng thành công một hệ thống tự động nhận dạng phong cách kiến "
           "trúc từ ảnh chụp theo hướng mở từ vựng, kết hợp nhiều mô hình thị giác – ngôn "
           "ngữ theo cơ chế hội đồng – trọng tài, có nối khớp cơ sở tri thức và cơ chế từ "
           "chối. Kết quả thử nghiệm cho thấy hệ thống đạt độ chính xác ngang bằng hoặc "
           "vượt các mô hình nền trên cùng thang đo, đồng thời mang lại các giá trị về khả "
           "năng mở rộng, độ tin cậy và khả năng giải thích. Đây là tiền đề tốt cho các "
           "phát triển và ứng dụng thực tế trong tương lai.")


# ════════════════════════════ TÀI LIỆU THAM KHẢO ══════════════════════════════
_REFERENCES = [
    "Ching, F. D. K., Jarzombek, M. M., & Prakash, V. (2017). A global history of "
    "architecture (3rd ed.). Wiley.",
    "Du, Y., Li, S., Torralba, A., Tenenbaum, J. B., & Mordatch, I. (2024). Improving "
    "factuality and reasoning in language models through multiagent debate. In "
    "Proceedings of the 41st International Conference on Machine Learning (ICML).",
    "He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image "
    "recognition. In Proceedings of the IEEE Conference on Computer Vision and Pattern "
    "Recognition (CVPR) (pp. 770–778).",
    "The J. Paul Getty Trust. (n.d.). Art & Architecture Thesaurus (AAT). Retrieved "
    "from https://www.getty.edu/research/tools/vocabularies/aat/",
    "Radford, A., Kim, J. W., Hallacy, C., Ramesh, A., Goh, G., Agarwal, S., … "
    "Sutskever, I. (2021). Learning transferable visual models from natural language "
    "supervision. In Proceedings of the 38th International Conference on Machine "
    "Learning (ICML) (pp. 8748–8763).",
    "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., "
    "Kaiser, Ł., & Polosukhin, I. (2017). Attention is all you need. In Advances in "
    "Neural Information Processing Systems (NeurIPS) (pp. 5998–6008).",
    "Wang, J., Wang, J., Athiwaratkun, B., Zhang, C., & Zou, J. (2024). Mixture-of-"
    "Agents enhances large language model capabilities. arXiv preprint arXiv:2406.04692.",
    "Wikidata contributors. (n.d.). Wikidata: The free knowledge base. Wikimedia "
    "Foundation. Retrieved from https://www.wikidata.org/",
    "Xu, Z., Tao, D., Zhang, Y., Wu, J., & Tsoi, A. C. (2014). Architectural style "
    "classification using multinomial latent logistic regression. In Computer Vision – "
    "ECCV 2014 (pp. 600–615). Springer.",
    "Zhang, S., Ai, W., Sun, Y., & Chen, T. (2025). ArchiLense: A framework for "
    "quantitative analysis of architectural styles based on vision large language "
    "models. arXiv preprint arXiv:2506.07739.",
]


def tai_lieu_tham_khao(d):
    p = d.d.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TÀI LIỆU THAM KHẢO")
    _set_font(r, size=SZ_CHAPTER, bold=True)
    for ref in _REFERENCES:
        para = d.d.add_paragraph(style="Normal")
        run = para.add_run(ref)
        _set_font(run)
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-1.0)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def main():
    if not os.path.exists(TEMPLATE):
        shutil.copy(SOURCE, TEMPLATE)
        print(f"Đã tạo bản sao template: {TEMPLATE}")
    document = docx.Document(TEMPLATE)

    _ensure_style(document, "CaptionHinh", align=WD_ALIGN_PARAGRAPH.CENTER)
    _ensure_style(document, "CaptionBang", align=WD_ALIGN_PARAGRAPH.LEFT)

    fix_global_formatting(document)
    rewrite_frontmatter(document)
    build_frontmatter_tail(document)
    clear_body_after_frontmatter(document)
    setup_page_numbering(document)

    doc = Doc(document)
    build_chapters(doc)

    try:
        document.save(OUTPUT)
        print(f"Đã lưu: {OUTPUT}")
    except PermissionError:
        fallback = OUTPUT.replace(".docx", ".new.docx")
        document.save(fallback)
        print(f"[!] {OUTPUT} đang bị khóa (mở trong Word). Đã lưu tạm: {fallback}")
        print("    → Đóng Word (KHÔNG lưu), xóa file cũ và đổi tên file .new.docx, "
              "hoặc chạy lại script.")


if __name__ == "__main__":
    main()
